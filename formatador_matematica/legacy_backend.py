
from __future__ import annotations

import json
import os
import re
import sys
import threading
import traceback
import subprocess
from dataclasses import dataclass, field, asdict
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

# ----------------------------
# Importações com tratamento
# ----------------------------
IMPORT_ERRORS = []

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk
except Exception as e:
    IMPORT_ERRORS.append(f"Tkinter: {e}")

try:
    from PIL import Image
except Exception as e:
    IMPORT_ERRORS.append(f"Pillow (PIL): {e}")

try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except Exception as e:
    IMPORT_ERRORS.append(f"python-docx: {e}")

try:
    import pythoncom
except Exception as e:
    IMPORT_ERRORS.append(f"pythoncom / pywin32: {e}")

try:
    import win32com.client as win32
except Exception as e:
    IMPORT_ERRORS.append(f"win32com / pywin32: {e}")

try:
    import fitz  # PyMuPDF
except Exception as e:
    IMPORT_ERRORS.append(f"PyMuPDF (fitz): {e}")

DND_AVAILABLE = False
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD  # type: ignore
    DND_AVAILABLE = True
except Exception:
    DND_AVAILABLE = False

WD_FORMAT_PDF = 17


def show_fatal_error_and_wait(msg: str):
    try:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("Erro ao iniciar o programa", msg)
        root.destroy()
    except Exception:
        pass

    print("\nERRO AO INICIAR O PROGRAMA\n")
    print(msg)
    print("\nPressione ENTER para fechar...")
    try:
        input()
    except Exception:
        pass
    sys.exit(1)


if IMPORT_ERRORS:
    show_fatal_error_and_wait(
        "Algumas dependências não estão instaladas corretamente:\n\n- "
        + "\n- ".join(IMPORT_ERRORS)
        + "\n\nInstale com:\n"
          "pip install pywin32 pymupdf pillow python-docx\n\n"
          "Opcional para arrastar e soltar:\n"
          "pip install tkinterdnd2"
    )


# ======================================================================
# BLOCO 1 - GERADOR SAFA
# ======================================================================

DEFAULT_FONT = "Arial"
DEFAULT_FONT_SIZE = 14
PAGE_WIDTH_CM = 19.0
PAGE_HEIGHT_CM = 25.4
MARGIN_CM = 1.27
IMAGE_WIDTH_CM = 15.5
ALT_BOX_WIDTH_CM = 16.0
EXTRACT_EMBEDDED_IMAGES = True

IGNORE_TAGS = (
    "MINI RELATÓRIO FINAL",
    "TABELA DE CHECAGEM TÉCNICA",
    "Suporte",
    "Gabarito",
    "Justificativa técnica dos distratores",
)


@dataclass
class ValidationIssue:
    level: str
    message: str


@dataclass
class ItemData:
    item_index: int
    ensinart: bool = True
    texto_base: str = ""
    enunciado: str = ""
    comando: str = ""
    alternativas: Dict[str, str] = field(default_factory=dict)
    explicacao_estudante: Dict[str, str] = field(default_factory=dict)
    gabarito: str = ""
    imagem_enunciado_path: str = ""
    imagens_alternativas: Dict[str, str] = field(default_factory=dict)
    issues: List[ValidationIssue] = field(default_factory=list)

    def to_json_dict(self) -> dict:
        data = asdict(self)
        data["issues"] = [asdict(x) for x in self.issues]
        return data


@dataclass
class RawBlock:
    kind: str
    text: str = ""
    image_bytes: Optional[bytes] = None
    image_name: str = ""


def normalize_spaces(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def append_text(current: str, new_text: str) -> str:
    if not new_text:
        return current.strip()
    if not current:
        return new_text.strip()
    return f"{current}\n{new_text}".strip()


def clean_alt_prefix(text: str) -> str:
    return re.sub(r"^[A-D]\)\s*", "", text.strip(), flags=re.IGNORECASE)


def clean_student_expl_prefix(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^(Correta|Incorreta)\.?\s*", "", text, flags=re.IGNORECASE)
    return text.strip()


def pad_alt_line(text: str, target_len: int = 110) -> str:
    text = text.strip()
    if len(text) >= target_len:
        return text
    return text + (" " * (target_len - len(text)))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def iter_block_items(parent) -> List[Tuple[str, object]]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("parent deve ser Document ou _Cell")

    items = []
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            items.append(("paragraph", Paragraph(child, parent)))
        elif child.tag.endswith("}tbl"):
            items.append(("table", Table(child, parent)))
    return items


def get_paragraph_image_bytes(paragraph: Paragraph) -> List[Tuple[str, bytes]]:
    images: List[Tuple[str, bytes]] = []
    rel_ids = []
    for blip in paragraph._element.xpath('.//a:blip'):
        embed = blip.get(qn('r:embed'))
        if embed:
            rel_ids.append(embed)
    for rel_id in rel_ids:
        part = paragraph.part.related_parts.get(rel_id)
        if part is not None:
            images.append((getattr(part, "filename", rel_id), part.blob))
    return images


def read_docx_blocks(path: Path) -> List[RawBlock]:
    doc = Document(path)
    blocks: List[RawBlock] = []

    for kind, obj in iter_block_items(doc):
        if kind == "paragraph":
            text = normalize_spaces(obj.text)
            if text:
                blocks.append(RawBlock(kind="paragraph", text=text))
            if EXTRACT_EMBEDDED_IMAGES:
                for idx, (name, blob) in enumerate(get_paragraph_image_bytes(obj)):
                    blocks.append(RawBlock(kind="image", image_bytes=blob, image_name=name or f"img_{idx}"))
        elif kind == "table":
            lines = []
            for row in obj.rows:
                cells = [normalize_spaces(c.text) for c in row.cells if normalize_spaces(c.text)]
                if cells:
                    lines.append(" | ".join(cells))
            if lines:
                blocks.append(RawBlock(kind="table", text="\n".join(lines)))
    return blocks


def split_items(blocks: List[RawBlock]) -> List[List[RawBlock]]:
    start_idx = None
    for i, block in enumerate(blocks):
        if block.kind == "paragraph" and re.search(r"\bITEM FINAL\b", block.text, re.IGNORECASE):
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("Nenhum 'ITEM FINAL' encontrado no arquivo.")

    blocks = blocks[start_idx:]
    items: List[List[RawBlock]] = []
    current: List[RawBlock] = []
    collecting = False

    for block in blocks:
        text = block.text if block.kind == "paragraph" else ""

        if block.kind == "paragraph" and re.search(r"\bITEM FINAL\b", text, re.IGNORECASE):
            if current:
                items.append(current)
            current = [block]
            collecting = True
            continue

        if collecting and block.kind == "paragraph" and re.match(r"^Item\s+\d+\b", text, re.IGNORECASE):
            if current:
                items.append(current)
            current = []
            collecting = False
            continue

        if collecting:
            current.append(block)

    if current:
        items.append(current)
    return items


def parse_item_selection(text: str, total_items: int) -> Set[int]:
    text = (text or '').strip()
    if not text:
        return set(range(1, total_items + 1))

    selected: Set[int] = set()
    # Aceita vírgula (",") ou ponto e vírgula (";") como separadores
    normalized = text.replace(';', ',')
    parts = [p.strip() for p in normalized.split(',') if p.strip()]
    if not parts:
        raise ValueError("Informe os itens no formato 3 ou 1,3,5-7.")

    for part in parts:
        if '-' in part:
            start_str, end_str = [x.strip() for x in part.split('-', 1)]
            if not start_str.isdigit() or not end_str.isdigit():
                raise ValueError(f"Faixa inválida: '{part}'. Use, por exemplo, 2-5.")
            start = int(start_str)
            end = int(end_str)
            if start > end:
                raise ValueError(f"Faixa inválida: '{part}'. O início deve ser menor ou igual ao fim.")
            for num in range(start, end + 1):
                if num < 1 or num > total_items:
                    raise ValueError(f"Item fora do intervalo disponível: {num}. Total no arquivo: {total_items}.")
                selected.add(num)
        else:
            if not part.isdigit():
                raise ValueError(f"Item inválido: '{part}'. Use números separados por vírgula.")
            num = int(part)
            if num < 1 or num > total_items:
                raise ValueError(f"Item fora do intervalo disponível: {num}. Total no arquivo: {total_items}.")
            selected.add(num)

    if not selected:
        raise ValueError("Nenhum item válido foi informado.")

    return selected


def save_image_bytes(blob: bytes, path: Path) -> Optional[Path]:
    ensure_dir(path.parent)
    try:
        image = Image.open(BytesIO(blob))
        image.save(path)
        return path
    except Exception:
        try:
            path.write_bytes(blob)
            return path
        except Exception:
            return None


def parse_item(blocks: List[RawBlock], item_index: int, image_dir: Path) -> ItemData:
    item = ItemData(item_index=item_index)

    image_paths: List[str] = []
    for img_idx, block in enumerate(blocks):
        if block.kind == "image" and block.image_bytes:
            out = image_dir / f"item_{item_index:02d}_img_{img_idx+1}.png"
            saved = save_image_bytes(block.image_bytes, out)
            if saved:
                image_paths.append(str(saved))

    mode = None
    expected_alt_idx = 0
    alt_order = ["A", "B", "C", "D"]
    expected_expl_idx = 0
    img_idx = 0

    for block in blocks:
        if block.kind == "paragraph":
            line = normalize_spaces(block.text)

            gabarito_match = re.match(r"^Gabarito:\s*([A-D])", line, re.IGNORECASE)
            if gabarito_match:
                item.gabarito = gabarito_match.group(1).upper()
                continue
            if not line:
                continue

            if re.fullmatch(r"ITEM FINAL:?", line, re.IGNORECASE):
                continue

            if any(re.match(rf"^{re.escape(tag)}\b", line, re.IGNORECASE) for tag in IGNORE_TAGS):
                mode = "ignore"
                continue

            if re.match(r"^Texto[\s\-]base\b", line, re.IGNORECASE):
                mode = "texto_base"
                content = re.sub(r"^Texto[\s\-]base\s*:?\s*", "", line, flags=re.IGNORECASE).strip()
                if content:
                    item.texto_base = append_text(item.texto_base, content)
                continue

            if re.match(r"^Enunciado\b", line, re.IGNORECASE):
                mode = "enunciado"
                content = re.sub(r"^Enunciado\s*:?\s*", "", line, flags=re.IGNORECASE).strip()
                if content:
                    item.enunciado = append_text(item.enunciado, content)
                continue

            if re.match(r"^Comando\b", line, re.IGNORECASE):
                mode = "comando"
                content = re.sub(r"^Comando\s*:?\s*", "", line, flags=re.IGNORECASE).strip()
                if content:
                    item.comando = append_text(item.comando, content)
                continue

            if re.match(r"^Alternativas\b", line, re.IGNORECASE):
                mode = "alternativas"
                expected_alt_idx = 0
                continue

            if re.match(r"^Explica[cç][aã]o ao (estudante|aluno)\b", line, re.IGNORECASE):
                mode = "explicacao"
                expected_expl_idx = 0
                continue

            alt_match = re.match(r"^([A-D])\)\s*(.*)$", line, re.IGNORECASE)
            if alt_match:
                letra = alt_match.group(1).upper()
                content = alt_match.group(2).strip()
                if mode == "alternativas":
                    # Sempre registra a alternativa, mesmo sem texto (pode ter só imagem)
                    if letra not in item.alternativas:
                        item.alternativas[letra] = content
                    else:
                        item.alternativas[letra] = (item.alternativas[letra] + " " + content).strip()
                elif mode == "explicacao":
                    item.explicacao_estudante[letra] = content
                continue

            # Detecta linhas que marcam fim de seção (usadas no fallback posicional)
            _is_section_break = bool(
                re.match(r"^Gabarito\b", line, re.IGNORECASE)
                or re.match(r"^Justificativa", line, re.IGNORECASE)
                or re.match(r"^Explica[cç][aã]o", line, re.IGNORECASE)
            )

            if mode == "texto_base":
                item.texto_base = append_text(item.texto_base, line)
            elif mode == "enunciado":
                item.enunciado = append_text(item.enunciado, line)
            elif mode == "comando":
                item.comando = append_text(item.comando, line)
            elif mode == "alternativas" and expected_alt_idx < 4:
                # Modo posicional: parar se encontrar início de outra seção conhecida
                if _is_section_break:
                    mode = None
                else:
                    letra = alt_order[expected_alt_idx]
                    item.alternativas[letra] = line
                    expected_alt_idx += 1
            elif mode == "explicacao" and expected_expl_idx < 4:
                letra = alt_order[expected_expl_idx]
                item.explicacao_estudante[letra] = line
                expected_expl_idx += 1
            elif mode == "explicacao" and item.explicacao_estudante:
                last = list(item.explicacao_estudante.keys())[-1]
                item.explicacao_estudante[last] = append_text(item.explicacao_estudante[last], line)

        elif block.kind == "image":
            if img_idx >= len(image_paths):
                continue
            current_img = image_paths[img_idx]
            img_idx += 1

            if mode in ("texto_base", "enunciado") and not item.imagem_enunciado_path:
                item.imagem_enunciado_path = current_img
            elif mode == "alternativas":
                # Atribui a imagem à próxima alternativa sem imagem ainda
                # (mesmo que o texto da alternativa esteja vazio)
                for key in alt_order:
                    if key not in item.imagens_alternativas:
                        # Garante que a alternativa existe mesmo sem texto
                        if key not in item.alternativas:
                            item.alternativas[key] = ""
                        item.imagens_alternativas[key] = current_img
                        break

    return item


def validate_item(item: ItemData) -> ItemData:
    if not item.enunciado:
        item.issues.append(ValidationIssue("error", "Enunciado ausente."))
    if not item.comando:
        item.issues.append(ValidationIssue("error", "Comando ausente."))
    for alt in ["A", "B", "C", "D"]:
        tem_texto = alt in item.alternativas
        tem_imagem = alt in item.imagens_alternativas
        if not tem_texto and not tem_imagem:
            item.issues.append(ValidationIssue("error", f"Alternativa {alt} ausente."))
        if alt not in item.explicacao_estudante:
            item.issues.append(ValidationIssue("warning", f"Explicação ao estudante da alternativa {alt} ausente."))
    return item


def infer_correct_alt(item: ItemData) -> str:
    if getattr(item, "gabarito", ""):
        return item.gabarito

    for alt in ["A", "B", "C", "D"]:
        text = item.explicacao_estudante.get(alt, "")
        if re.match(r"^\s*Correta\b", text, re.IGNORECASE):
            return alt
    return "A"


def build_response_lines(item: ItemData) -> List[Tuple[str, str]]:
    lines = []
    correct_alt = infer_correct_alt(item)

    for alt in ["A", "B", "C", "D"]:
        alt_text = item.alternativas.get(alt, "").strip()
        raw_expl = item.explicacao_estudante.get(alt, "").strip()

        explanation = re.sub(r"^(Correta|Incorreta)\.?\s*", "", raw_expl, flags=re.IGNORECASE).strip()

        status = "ESTÁ CORRETA." if alt == correct_alt else "ESTÁ INCORRETA."

        # Se a alternativa é só imagem (sem texto), usa rótulo descritivo
        if alt_text:
            titulo = f'"{alt_text}" {status}'
        else:
            tem_imagem = bool(item.imagens_alternativas.get(alt))
            rotulo = "[IMAGEM]" if tem_imagem else "[SEM TEXTO]"
            titulo = f'Alternativa {alt} {rotulo} {status}'

        corpo = explanation

        lines.append((titulo, corpo))

    return lines


class StandardDocxBuilder:
    def build(self, item: ItemData, output_path: Path) -> None:
        doc = Document()
        self._setup_document(doc)
        self._add_page_1(doc, item)

        for alt in ["A", "B", "C", "D"]:
            doc.add_page_break()
            self._add_alternative_page(doc, item.alternativas.get(alt, ""), item.imagens_alternativas.get(alt, ""))

        doc.add_page_break()
        self._add_response_page(doc, item)

        ensure_dir(output_path.parent)
        doc.save(str(output_path))

    def _setup_document(self, doc: _Document) -> None:
        section = doc.sections[0]
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)

        style = doc.styles["Normal"]
        style.font.name = DEFAULT_FONT
        style.font.size = Pt(DEFAULT_FONT_SIZE)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
        # Forçar justificacao como padrao do estilo Normal
        pPr = style._element.get_or_add_pPr()
        for old in pPr.findall(qn("w:jc")):
            pPr.remove(old)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

    def _style_run(self, run, size_pt: int = DEFAULT_FONT_SIZE):
        run.font.name = DEFAULT_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
        run.font.size = Pt(size_pt)

    def _add_paragraph(self, doc: _Document, text: str, justify: bool = True):
        p = doc.add_paragraph()
        lines = text.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            self._style_run(run)
            if i < len(lines) - 1:
                run.add_break()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        return p

    def _add_picture(self, doc: _Document, image_path: str):
        if not image_path:
            return
        p = Path(image_path)
        if not p.exists():
            return
        try:
            doc.add_picture(str(p), width=Cm(IMAGE_WIDTH_CM))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            self._add_paragraph(doc, f"[IMAGEM NÃO INSERIDA: {image_path}]", justify=False)

    def _add_page_1(self, doc: _Document, item: ItemData):
        if item.ensinart:
            self._add_paragraph(doc, "(ENSINART)", justify=False)
        if item.texto_base:
            self._add_paragraph(doc, item.texto_base)
        if item.enunciado:
            self._add_paragraph(doc, item.enunciado)
        self._add_picture(doc, item.imagem_enunciado_path)
        if item.comando:
            self._add_paragraph(doc, item.comando)

    def _add_alternative_page(self, doc: _Document, alt_text: str, image_path: str):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Cm(ALT_BOX_WIDTH_CM)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(pad_alt_line(clean_alt_prefix(alt_text)))
        self._style_run(run)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        if image_path:
            self._add_picture(doc, image_path)

    def _add_response_page(self, doc: _Document, item: ItemData):
        for titulo, explicacao in build_response_lines(item):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(titulo + " ")
            self._style_run(run)
            run.bold = True

            if explicacao:
                run2 = p.add_run("→ " + explicacao)
                self._style_run(run2)

            doc.add_paragraph("")


class SafaConverter:
    def __init__(self):
        self.builder = StandardDocxBuilder()

    def process_file(self, input_docx: Path, output_root: Path, ensinart: bool = True, logger=None, item_selection: str = "") -> dict:
        ensure_dir(output_root)
        blocks = read_docx_blocks(input_docx)
        items = split_items(blocks)
        selected_items = parse_item_selection(item_selection, len(items))

        report = {
            "arquivo_fonte": str(input_docx),
            "total_itens": len(items),
            "itens_selecionados": sorted(selected_items),
            "total_itens_selecionados": len(selected_items),
            "itens_ok": 0,
            "itens_com_erro": 0,
            "itens_ignorados": len(items) - len(selected_items),
            "itens": [],
        }

        for idx, block in enumerate(items, start=1):
            if idx not in selected_items:
                if logger:
                    logger(f"  - Item {idx:02d} ignorado (fora da seleção)")
                continue

            item = parse_item(block, idx, output_root / "_imagens_extraidas")
            item.ensinart = ensinart
            item = validate_item(item)

            output_docx = output_root / f"item_{idx:02d}_padronizado.docx"

            has_error = any(x.level == "error" for x in item.issues)
            if has_error:
                report["itens_com_erro"] += 1
                status = "erro"
                if logger:
                    logger(f"  - Item {idx:02d} com erro estrutural")
            else:
                self.builder.build(item, output_docx)
                report["itens_ok"] += 1
                status = "ok"
                if logger:
                    logger(f"  - Item {idx:02d} gerado")

            report["itens"].append({
                "item": idx,
                "status": status,
                "saida_docx": str(output_docx),
                "mensagens": [f"[{x.level.upper()}] {x.message}" for x in item.issues],
            })

        (output_root / "relatorio_processamento.json").write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        return report


class SafaTab(ttk.Frame):
    def __init__(self, master):
        super().__init__(master, padding=14)
        self.files: List[Path] = []
        self.output_dir = tk.StringVar()
        self.ensinart_var = tk.BooleanVar(value=True)
        self.status_var = tk.StringVar(value="Pronto")
        self.process_mode_var = tk.StringVar(value="todos")
        self.item_selection_var = tk.StringVar()
        self._build_ui()

    def _build_ui(self):
        ttk.Label(self, text="Gerador SAFA — Processamento em lote", style="Title.TLabel").pack(anchor="w", pady=(0, 12))

        frame_files = ttk.LabelFrame(self, text="1. Arquivos de entrada", style="Card.TLabelframe", padding=10)
        frame_files.pack(fill="x", pady=(0, 10))
        row_btns = ttk.Frame(frame_files)
        row_btns.pack(fill="x", pady=(0, 8))
        ttk.Button(row_btns, text="➕  Adicionar arquivos .docx", command=self.add_files, style="Accent.TButton").pack(side="left")
        ttk.Button(row_btns, text="➖  Remover selecionado", command=self.remove_selected, style="Danger.TButton").pack(side="left", padx=8)
        ttk.Button(row_btns, text="🗑️  Limpar lista", command=self.clear_files, style="Soft.TButton").pack(side="left")
        self.listbox = tk.Listbox(frame_files, height=10, selectmode=tk.EXTENDED)
        self.listbox.pack(fill="x")

        frame_out = ttk.LabelFrame(self, text="2. Pasta de saída", style="Card.TLabelframe", padding=10)
        frame_out.pack(fill="x", pady=(0, 10))
        row_out = ttk.Frame(frame_out)
        row_out.pack(fill="x")
        ttk.Entry(row_out, textvariable=self.output_dir).pack(side="left", fill="x", expand=True)
        ttk.Button(row_out, text="📁  Escolher pasta", command=self.choose_output_dir, style="Soft.TButton").pack(side="left", padx=(8, 0))

        frame_cfg = ttk.LabelFrame(self, text="3. Configurações", style="Card.TLabelframe", padding=10)
        frame_cfg.pack(fill="x", pady=(0, 10))
        ttk.Checkbutton(frame_cfg, text="Inserir linha (ENSINART) na primeira página", variable=self.ensinart_var).pack(anchor="w")

        frame_selection = ttk.Frame(frame_cfg)
        frame_selection.pack(fill="x", pady=(8, 0))
        ttk.Radiobutton(frame_selection, text="Processar todos os itens", value="todos", variable=self.process_mode_var, command=self._toggle_item_selection).pack(anchor="w")
        ttk.Radiobutton(frame_selection, text="Processar apenas itens específicos", value="selecionados", variable=self.process_mode_var, command=self._toggle_item_selection).pack(anchor="w", pady=(2, 0))

        row_items = ttk.Frame(frame_cfg)
        row_items.pack(fill="x", pady=(6, 0))
        ttk.Label(row_items, text="Itens:").pack(side="left")
        self.entry_item_selection = ttk.Entry(row_items, textvariable=self.item_selection_var)
        self.entry_item_selection.pack(side="left", fill="x", expand=True, padx=(8, 0))

        ttk.Label(frame_cfg, text="Exemplos: 3  |  1,4,7  |  2-5  |  1,3,8-10  |  3; 5-9").pack(anchor="w", pady=(4, 0))
        ttk.Label(frame_cfg, text="Esta versão tenta capturar imagens embutidas do Word.").pack(anchor="w", pady=(6, 0))

        row_action = ttk.Frame(self)
        row_action.pack(fill="x", pady=(0, 10))
        self.btn_process = ttk.Button(row_action, text="▶  Processar arquivos", command=self.start, style="Success.TButton")
        self.btn_process.pack(side="left")
        ttk.Label(row_action, textvariable=self.status_var, style="Status.TLabel").pack(side="left", padx=12)

        frame_log = ttk.LabelFrame(self, text="4. Andamento", style="Card.TLabelframe", padding=10)
        frame_log.pack(fill="both", expand=True)
        self.log_text = tk.Text(frame_log, wrap="word", height=18)
        self.log_text.pack(side="left", fill="both", expand=True)
        scroll = ttk.Scrollbar(frame_log, orient="vertical", command=self.log_text.yview)
        scroll.pack(side="right", fill="y")
        self.log_text.config(yscrollcommand=scroll.set)
        self._toggle_item_selection()

    def _toggle_item_selection(self):
        state = "normal" if self.process_mode_var.get() == "selecionados" else "disabled"
        self.entry_item_selection.config(state=state)

    def add_files(self):
        paths = filedialog.askopenfilenames(title="Selecione os arquivos Word", filetypes=[("Arquivos Word", "*.docx")])
        for path in paths:
            p = Path(path)
            if p not in self.files:
                self.files.append(p)
                self.listbox.insert(tk.END, str(p))

    def remove_selected(self):
        for idx in reversed(list(self.listbox.curselection())):
            self.listbox.delete(idx)
            del self.files[idx]

    def clear_files(self):
        self.listbox.delete(0, tk.END)
        self.files.clear()

    def choose_output_dir(self):
        path = filedialog.askdirectory(title="Escolha a pasta de saída")
        if path:
            self.output_dir.set(path)

    def log(self, msg: str):
        self.log_text.insert(tk.END, msg + "\n")
        self.log_text.see(tk.END)
        self.update_idletasks()

    def start(self):
        if not self.files:
            messagebox.showwarning("Aviso", "Selecione pelo menos um arquivo .docx.")
            return
        if not self.output_dir.get().strip():
            messagebox.showwarning("Aviso", "Escolha a pasta de saída.")
            return
        if self.process_mode_var.get() == "selecionados" and not self.item_selection_var.get().strip():
            messagebox.showwarning("Aviso", "Informe os itens desejados. Ex.: 3 ou 1,4,7-9.")
            return

        self.btn_process.config(state="disabled")
        self.status_var.set("Processando...")
        self.log_text.delete("1.0", tk.END)
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        try:
            converter = SafaConverter()
            total_files = 0
            total_ok = 0
            total_selected = 0
            item_selection = self.item_selection_var.get().strip() if self.process_mode_var.get() == "selecionados" else ""

            for file_path in self.files:
                self.log(f"Lendo arquivo: {file_path.name}")
                out_dir = Path(self.output_dir.get().strip()) / file_path.stem
                report = converter.process_file(
                    file_path,
                    out_dir,
                    ensinart=self.ensinart_var.get(),
                    logger=self.log,
                    item_selection=item_selection,
                )
                total_files += 1
                total_ok += report["itens_ok"]
                total_selected += report["total_itens_selecionados"]

            self.after(0, lambda: self._finish_ok(total_files, total_ok, total_selected))
        except Exception as e:
            msg = str(e)
            self.after(0, lambda m=msg: self._finish_error(m))

    def _finish_ok(self, total_files: int, total_ok: int, total_selected: int):
        self.btn_process.config(state="normal")
        self.status_var.set("Concluído")
        self.log(f"\nFinalizado. Arquivos: {total_files} | Itens selecionados: {total_selected} | Itens gerados: {total_ok}")
        messagebox.showinfo(
            "Concluído",
            f"Processamento finalizado.\n\nArquivos: {total_files}\nItens selecionados: {total_selected}\nItens gerados: {total_ok}"
        )

    def _finish_error(self, msg: str):
        self.btn_process.config(state="normal")
        self.status_var.set("Erro")
        self.log(f"\nERRO: {msg}")
        messagebox.showerror("Erro", msg)


# ======================================================================
# BLOCO 2 - WORD PARA PNG
# ======================================================================

def crop_vertical_keep_full_width(
    image: Image.Image,
    threshold: int = 245,
    extra_margin_px: int = 12,
) -> Image.Image:
    if image.mode != "RGB":
        image = image.convert("RGB")

    width, height = image.size
    pixels = image.load()

    top = None
    bottom = None

    for y in range(height):
        for x in range(width):
            r, g, b = pixels[x, y]
            if not (r >= threshold and g >= threshold and b >= threshold):
                top = y
                break
        if top is not None:
            break

    if top is None:
        return image.copy()

    for y in range(height - 1, -1, -1):
        for x in range(width):
            r, g, b = pixels[x, y]
            if not (r >= threshold and g >= threshold and b >= threshold):
                bottom = y
                break
        if bottom is not None:
            break

    top = max(0, top - extra_margin_px)
    bottom = min(height - 1, bottom + extra_margin_px)

    return image.crop((0, top, width, bottom + 1))


def export_word_to_pdf(word_path: Path, pdf_path: Path, word_app) -> None:
    doc = None
    try:
        doc = word_app.Documents.Open(str(word_path), ReadOnly=False)
        # Forcar justificacao em todos os paragrafos antes de exportar para PDF
        WD_ALIGN_JUSTIFY = 3  # wdAlignParagraphJustify
        for para in doc.Paragraphs:
            try:
                para.Alignment = WD_ALIGN_JUSTIFY
            except Exception:
                pass
        doc.SaveAs2(str(pdf_path), FileFormat=WD_FORMAT_PDF)
    finally:
        if doc is not None:
            doc.Close(False)


def pdf_pages_to_pngs(
    pdf_path: Path,
    out_dir: Path,
    dpi: int = 220,
    crop_vertical: bool = True,
    threshold: int = 245,
    extra_margin_px: int = 12,
) -> List[Path]:
    generated = []
    doc = fitz.open(str(pdf_path))
    try:
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            if crop_vertical:
                img = crop_vertical_keep_full_width(
                    img,
                    threshold=threshold,
                    extra_margin_px=extra_margin_px,
                )

            output_file = out_dir / f"pagina_{i:02d}.png"
            img.save(output_file, "PNG")
            generated.append(output_file)
    finally:
        doc.close()

    return generated


def safe_folder_name(name: str) -> str:
    invalid = '<>:"/\\|?*'
    for ch in invalid:
        name = name.replace(ch, "_")
    return name.strip().rstrip(".")


def open_folder_in_explorer(path: Path):
    path = path.resolve()
    if os.name == "nt":
        os.startfile(str(path))
    elif sys.platform == "darwin":
        subprocess.run(["open", str(path)], check=False)
    else:
        subprocess.run(["xdg-open", str(path)], check=False)


def process_documents(
    word_files: List[Path],
    base_output_dir: Path,
    dpi: int,
    crop_vertical: bool,
    threshold: int,
    extra_margin_px: int,
    log_callback=None,
    progress_callback=None,
):
    pythoncom.CoInitialize()
    word_app = None
    success_count = 0

    try:
        try:
            word_app = win32.DispatchEx("Word.Application")
        except Exception as e:
            raise RuntimeError(
                "Não foi possível iniciar o Microsoft Word automaticamente. "
                "Verifique se o Word está instalado neste computador."
            ) from e

        word_app.Visible = False
        word_app.DisplayAlerts = 0

        total = len(word_files)

        for idx, word_file in enumerate(word_files, start=1):
            doc_name = word_file.stem
            final_out_dir = base_output_dir / safe_folder_name(doc_name)
            temp_out_dir = base_output_dir / f"__tmp__{safe_folder_name(doc_name)}"
            temp_pdf = base_output_dir / f"__tmp__{safe_folder_name(doc_name)}.pdf"

            if temp_out_dir.exists():
                for p in temp_out_dir.glob("*"):
                    if p.is_file():
                        p.unlink()
                try:
                    temp_out_dir.rmdir()
                except OSError:
                    pass

            if temp_pdf.exists():
                temp_pdf.unlink()

            if progress_callback:
                progress_callback(idx - 1, total, f"Convertendo: {word_file.name}")

            if log_callback:
                log_callback(f"[{idx}/{total}] {word_file.name}")

            try:
                temp_out_dir.mkdir(parents=True, exist_ok=True)

                export_word_to_pdf(word_file, temp_pdf, word_app)

                generated = pdf_pages_to_pngs(
                    temp_pdf,
                    temp_out_dir,
                    dpi=dpi,
                    crop_vertical=crop_vertical,
                    threshold=threshold,
                    extra_margin_px=extra_margin_px,
                )

                if not generated:
                    raise RuntimeError("Nenhuma imagem foi gerada.")

                if final_out_dir.exists():
                    for p in final_out_dir.glob("*"):
                        if p.is_file():
                            p.unlink()
                    try:
                        final_out_dir.rmdir()
                    except OSError:
                        pass

                temp_out_dir.rename(final_out_dir)
                success_count += 1

                if log_callback:
                    log_callback(f"    OK -> {len(generated)} página(s)")
                    log_callback(f"    Pasta: {final_out_dir}")

            except Exception as e:
                if temp_out_dir.exists():
                    for p in temp_out_dir.glob("*"):
                        if p.is_file():
                            p.unlink()
                    try:
                        temp_out_dir.rmdir()
                    except OSError:
                        pass

                if log_callback:
                    log_callback(f"    ERRO -> {e}")

            finally:
                if temp_pdf.exists():
                    temp_pdf.unlink()

            if progress_callback:
                progress_callback(idx, total, f"Finalizado: {word_file.name}")

    finally:
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    return success_count


class WordPngTab(ttk.Frame):
    def __init__(self, root):
        super().__init__(root, padding=14)
        self.root = root
        self.files: List[Path] = []
        self.last_output_dir: Optional[Path] = None

        self.output_dir_var = tk.StringVar()
        self.dpi_var = tk.IntVar(value=220)
        self.crop_var = tk.BooleanVar(value=True)
        self.threshold_var = tk.IntVar(value=245)
        self.margin_var = tk.IntVar(value=12)
        self.status_var = tk.StringVar(value="Pronto para começar.")
        self.progress_var = tk.DoubleVar(value=0)

        self.build_ui()

    def _colors(self):
        return _CURRENT_THEME if _CURRENT_THEME else THEMES["light"]

    def build_ui(self):
        top_cards = ttk.Frame(self)
        top_cards.pack(fill="x")

        self.build_files_card(top_cards)
        self.build_output_card(top_cards)

        footer = ttk.Frame(self)
        footer.pack(fill="x", pady=(8, 0))

        ttk.Progressbar(
            footer,
            variable=self.progress_var,
            maximum=100,
            style="Horizontal.TProgressbar"
        ).pack(side="top", fill="x", expand=True, pady=(0, 6))

        footer_bottom = ttk.Frame(footer)
        footer_bottom.pack(fill="x")

        ttk.Label(footer_bottom, textvariable=self.status_var, style="Status.TLabel").pack(side="left")

        right = ttk.Frame(footer_bottom)
        right.pack(side="right")

        ttk.Button(right, text="📂  Abrir pasta final", command=self.open_last_folder, style="Soft.TButton").pack(side="right", padx=(8, 0))
        ttk.Button(right, text="▶  Gerar PNGs", command=self.run, style="Success.TButton").pack(side="right", padx=(8, 0))
        ttk.Button(right, text="🧹  Limpar tudo", command=self.reset_all, style="Soft.TButton").pack(side="right")

        bottom_cards = ttk.Frame(self)
        bottom_cards.pack(fill="both", expand=True, pady=(8, 0))

        self.build_config_card(bottom_cards)
        self.build_log_card(bottom_cards)

    def build_files_card(self, parent):
        card = ttk.LabelFrame(parent, text="1. Arquivos do Word", style="Card.TLabelframe", padding=12)
        card.pack(side="left", fill="both", expand=True, padx=(0, 6))

        inner = ttk.Frame(card, style="CardInner.TFrame")
        inner.pack(fill="both", expand=True)

        row = ttk.Frame(inner, style="CardInner.TFrame")
        row.pack(fill="x", pady=(0, 10))

        ttk.Button(row, text="➕  Adicionar arquivos", command=self.add_files, style="Accent.TButton").pack(side="left")
        ttk.Button(row, text="➖  Remover selecionado", command=self.remove_selected, style="Danger.TButton").pack(side="left", padx=6)
        ttk.Button(row, text="🗑️  Limpar lista", command=self.clear_files, style="Soft.TButton").pack(side="left")

        hint = "Arraste arquivos .doc ou .docx para a área abaixo."
        if not DND_AVAILABLE:
            hint += "  (Arrastar e soltar opcional: instale tkinterdnd2)"
        ttk.Label(inner, text=hint, style="Mini.TLabel").pack(anchor="w", pady=(0, 8))

        drop_container = tk.Frame(inner, bg=self._colors()["drop_bg"], highlightbackground=self._colors()["drop_border"], highlightthickness=1)
        drop_container.pack(fill="both", expand=True)

        c = self._colors()
        self.listbox = tk.Listbox(
            drop_container,
            height=12,
            activestyle="none",
            selectmode=tk.EXTENDED,
            font=("Segoe UI", 10),
            bd=0,
            highlightthickness=0,
            relief="flat",
            background=c["drop_bg"],
            foreground=c["fg"],
            selectbackground=c["select_bg"],
            selectforeground=c["select_fg"],
        )
        self.listbox.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        scrollbar = ttk.Scrollbar(drop_container, orient="vertical", command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y", pady=10, padx=(0, 8))
        self.listbox.configure(yscrollcommand=scrollbar.set)

        if DND_AVAILABLE:
            self.listbox.drop_target_register(DND_FILES)
            self.listbox.dnd_bind("<<Drop>>", self.on_drop_files)
            drop_container.drop_target_register(DND_FILES)
            drop_container.dnd_bind("<<Drop>>", self.on_drop_files)

        self.files_info = ttk.Label(inner, text="Nenhum arquivo selecionado.", style="Mini.TLabel")
        self.files_info.pack(anchor="w", pady=(8, 0))

    def build_output_card(self, parent):
        card = ttk.LabelFrame(parent, text="2. Pasta de saída", style="Card.TLabelframe", padding=12)
        card.pack(side="left", fill="both", expand=True, padx=(6, 0))

        inner = ttk.Frame(card, style="CardInner.TFrame")
        inner.pack(fill="both", expand=True)

        ttk.Label(inner, text="Escolha a pasta onde o programa criará as subpastas de cada documento.", style="Mini.TLabel").pack(anchor="w", pady=(0, 10))

        row = ttk.Frame(inner, style="CardInner.TFrame")
        row.pack(fill="x")

        ttk.Entry(row, textvariable=self.output_dir_var, font=("Segoe UI", 10)).pack(side="left", fill="x", expand=True, padx=(0, 8), ipady=4)
        ttk.Button(row, text="📁  Escolher pasta", command=self.choose_output_dir, style="Soft.TButton").pack(side="left")

        help_box = tk.Text(
            inner,
            height=8,
            wrap="word",
            bd=0,
            highlightthickness=0,
            font=("Segoe UI", 10),
            foreground=self._colors()["muted"],
            background=self._colors()["card"],
        )
        help_box.pack(fill="both", expand=True, pady=(12, 0))
        help_box.insert(
            "1.0",
            "Saída:\n\n"
            "• Um documento gera uma pasta própria.\n"
            "• Uma página gera um PNG.\n"
            "• A largura da imagem continua sendo a largura total da folha.\n"
            "• O corte remove apenas topo e rodapé em branco."
        )
        help_box.configure(state="disabled")

    def build_config_card(self, parent):
        card = ttk.LabelFrame(parent, text="3. Configurações", style="Card.TLabelframe", padding=12)
        card.pack(side="left", fill="both", expand=False, padx=(0, 6), ipadx=10)

        inner = ttk.Frame(card, style="CardInner.TFrame")
        inner.pack(fill="both", expand=True)

        ttk.Checkbutton(
            inner,
            text="✂ Cortar só a altura, mantendo a largura total da folha",
            variable=self.crop_var,
        ).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 12))

        ttk.Label(inner, text="DPI das imagens", style="Mini.TLabel").grid(row=1, column=0, sticky="w")
        ttk.Spinbox(inner, from_=72, to=600, textvariable=self.dpi_var, width=10).grid(row=2, column=0, sticky="w", pady=(2, 12))

        ttk.Label(inner, text="Limite de branco", style="Mini.TLabel").grid(row=3, column=0, sticky="w")
        ttk.Spinbox(inner, from_=200, to=254, textvariable=self.threshold_var, width=10).grid(row=4, column=0, sticky="w", pady=(2, 12))

        ttk.Label(inner, text="Margem extra vertical (px)", style="Mini.TLabel").grid(row=5, column=0, sticky="w")
        ttk.Spinbox(inner, from_=0, to=500, textvariable=self.margin_var, width=10).grid(row=6, column=0, sticky="w", pady=(2, 12))

        tips = tk.Text(
            inner,
            height=10,
            wrap="word",
            bd=0,
            highlightthickness=0,
            font=("Segoe UI", 10),
            foreground=self._colors()["muted"],
            background=self._colors()["card"],
        )
        tips.grid(row=7, column=0, sticky="nsew", pady=(8, 0))
        tips.insert(
            "1.0",
            "Sugestões:\n\n"
            "• DPI 220: bom equilíbrio entre qualidade e tamanho.\n"
            "• Limite 245: funciona bem em páginas brancas.\n"
            "• Margem 12 px: pequeno respiro acima e abaixo do conteúdo.\n"
            "• Se o corte ficar agressivo, reduza o limite de branco."
        )
        tips.configure(state="disabled")
        inner.rowconfigure(7, weight=1)

    def build_log_card(self, parent):
        card = ttk.LabelFrame(parent, text="4. Andamento", style="Card.TLabelframe", padding=12)
        card.pack(side="left", fill="both", expand=True, padx=(6, 0))

        inner = ttk.Frame(card, style="CardInner.TFrame")
        inner.pack(fill="both", expand=True)

        self.log = tk.Text(
            inner,
            wrap="word",
            font=("Consolas", 10),
            bd=0,
            highlightthickness=1,
            relief="flat",
            background=self._colors()["input_bg"],
            foreground=self._colors()["fg"],
        )
        self.log.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(inner, orient="vertical", command=self.log.yview)
        scrollbar.pack(side="right", fill="y")
        self.log.configure(yscrollcommand=scrollbar.set)

    def normalize_dropped_files(self, raw_data: str) -> List[Path]:
        results = []
        current = ""
        in_braces = False

        for ch in raw_data:
            if ch == "{":
                in_braces = True
                current = ""
            elif ch == "}":
                in_braces = False
                if current.strip():
                    results.append(Path(current.strip()))
                current = ""
            elif ch == " " and not in_braces:
                if current.strip():
                    results.append(Path(current.strip()))
                    current = ""
            else:
                current += ch

        if current.strip():
            results.append(Path(current.strip()))

        return results

    def on_drop_files(self, event):
        try:
            paths = self.normalize_dropped_files(event.data)
            added = 0
            for p in paths:
                if p.suffix.lower() in [".doc", ".docx"] and p.exists():
                    if p not in self.files:
                        self.files.append(p)
                        self.listbox.insert(tk.END, str(p))
                        added += 1
            self.update_files_info()
            if added:
                self.status_var.set(f"{added} arquivo(s) adicionados por arrastar e soltar.")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível adicionar os arquivos arrastados.\n\n{e}")

    def add_files(self):
        selected = filedialog.askopenfilenames(
            title="Selecione documentos Word",
            filetypes=[("Documentos Word", "*.docx *.doc")]
        )
        for item in selected:
            p = Path(item)
            if p not in self.files:
                self.files.append(p)
                self.listbox.insert(tk.END, str(p))
        self.update_files_info()

    def remove_selected(self):
        indices = list(self.listbox.curselection())
        for idx in reversed(indices):
            self.listbox.delete(idx)
            del self.files[idx]
        self.update_files_info()

    def clear_files(self):
        self.listbox.delete(0, tk.END)
        self.files.clear()
        self.update_files_info()

    def reset_all(self):
        self.clear_files()
        self.output_dir_var.set("")
        self.dpi_var.set(220)
        self.crop_var.set(True)
        self.threshold_var.set(245)
        self.margin_var.set(12)
        self.progress_var.set(0)
        self.last_output_dir = None
        self.status_var.set("Pronto para começar.")
        self.log.delete("1.0", tk.END)

    def update_files_info(self):
        qtd = len(self.files)
        if qtd == 0:
            self.files_info.config(text="Nenhum arquivo selecionado.")
        elif qtd == 1:
            self.files_info.config(text="1 arquivo selecionado.")
        else:
            self.files_info.config(text=f"{qtd} arquivos selecionados.")

    def choose_output_dir(self):
        folder = filedialog.askdirectory(title="Escolha a pasta de saída")
        if folder:
            self.output_dir_var.set(folder)

    def write_log(self, text: str):
        self.log.insert(tk.END, text + "\n")
        self.log.see(tk.END)
        self.update_idletasks()

    def set_progress(self, done: int, total: int, status: str):
        pct = 0 if total == 0 else (done / total) * 100
        self.progress_var.set(pct)
        self.status_var.set(status)
        self.update_idletasks()

    def open_last_folder(self):
        target = None
        if self.last_output_dir and self.last_output_dir.exists():
            target = self.last_output_dir
        elif self.output_dir_var.get().strip():
            p = Path(self.output_dir_var.get().strip())
            if p.exists():
                target = p

        if not target:
            messagebox.showinfo("Abrir pasta", "Ainda não há uma pasta final disponível para abrir.")
            return

        try:
            open_folder_in_explorer(target)
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir a pasta.\n\n{e}")

    def validate(self) -> Optional[str]:
        if not self.files:
            return "Adicione pelo menos um arquivo Word."
        if not self.output_dir_var.get().strip():
            return "Escolha a pasta de saída."
        try:
            dpi = int(self.dpi_var.get())
            if dpi < 72 or dpi > 600:
                return "O DPI deve ficar entre 72 e 600."
        except Exception:
            return "DPI inválido."
        try:
            threshold = int(self.threshold_var.get())
            if threshold < 200 or threshold > 254:
                return "O limite de branco deve ficar entre 200 e 254."
        except Exception:
            return "Limite de branco inválido."
        try:
            margin = int(self.margin_var.get())
            if margin < 0 or margin > 500:
                return "A margem extra deve ficar entre 0 e 500."
        except Exception:
            return "Margem extra inválida."
        return None

    def run(self):
        error = self.validate()
        if error:
            messagebox.showerror("Validação", error)
            return

        output_dir = Path(self.output_dir_var.get())
        output_dir.mkdir(parents=True, exist_ok=True)
        self.last_output_dir = output_dir

        self.log.delete("1.0", tk.END)
        self.write_log("Iniciando processamento...")
        self.set_progress(0, len(self.files), "Preparando conversão...")

        self.winfo_toplevel().config(cursor="watch")
        self.update_idletasks()

        # Captura os parâmetros antes de entrar na thread
        word_files = list(self.files)
        dpi = int(self.dpi_var.get())
        crop_vertical = bool(self.crop_var.get())
        threshold = int(self.threshold_var.get())
        extra_margin_px = int(self.margin_var.get())

        def _run_thread():
            try:
                success_count = process_documents(
                    word_files=word_files,
                    base_output_dir=output_dir,
                    dpi=dpi,
                    crop_vertical=crop_vertical,
                    threshold=threshold,
                    extra_margin_px=extra_margin_px,
                    log_callback=lambda msg: self.after(0, lambda m=msg: self.write_log(m)),
                    progress_callback=lambda done, total, status: self.after(
                        0, lambda d=done, t=total, s=status: self.set_progress(d, t, s)
                    ),
                )
                self.after(0, lambda sc=success_count: self._finish_png_ok(sc, len(word_files)))
            except Exception:
                err = traceback.format_exc()
                self.after(0, lambda e=err: self._finish_png_error(e))

        threading.Thread(target=_run_thread, daemon=True).start()

    def _finish_png_ok(self, success_count: int, total: int):
        self.set_progress(total, total, "Processamento concluído.")
        self.write_log("Concluído.")
        self.winfo_toplevel().config(cursor="")
        self.update_idletasks()

        if success_count > 0:
            if messagebox.askyesno(
                "Concluído",
                f"{success_count} documento(s) convertido(s) com sucesso.\n\nDeseja abrir a pasta final agora?"
            ):
                self.open_last_folder()
        else:
            messagebox.showwarning(
                "Concluído sem sucesso",
                "O processamento terminou, mas nenhum documento foi convertido com sucesso.\n\nVerifique o log."
            )

    def _finish_png_error(self, err: str):
        self.write_log("ERRO GERAL:")
        self.write_log(err)
        self.status_var.set("Ocorreu um erro durante o processamento.")
        self.winfo_toplevel().config(cursor="")
        self.update_idletasks()
        messagebox.showerror("Erro", "Ocorreu um erro durante o processamento.\n\nVeja o log para mais detalhes.")




# ======================================================================
# BLOCO 3 - FORMATADOR DE ITEM PARA WORD
# ======================================================================

ITEM_FORMATTER_TITLE = 'Formatador de Itens para Word'
ITEM_FORMATTER_DEFAULT_FILENAME = 'item_formatado.docx'


def itemfmt_set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def itemfmt_set_cell_border(cell, color='808080', size='6'):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = f'w:{edge}'
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def itemfmt_clean_line(line: str) -> str:
    return re.sub(r'\s+', ' ', line.strip())


def itemfmt_find_line_idx(lines, startswith_text):
    target = startswith_text.lower()
    for i, line in enumerate(lines):
        if itemfmt_clean_line(line).lower().startswith(target):
            return i
    return -1


def itemfmt_parse_bullets(lines):
    bullets = []
    for line in lines:
        clean = itemfmt_clean_line(line)
        if not clean:
            continue
        if clean.startswith('·') or clean.startswith('-') or clean.startswith('•'):
            bullets.append(re.sub(r'^[·•\-]\s*', '', clean).strip())
        elif ':' in clean or clean.lower().startswith('matriz-base') or clean.lower().startswith('ano-alvo') or clean.lower().startswith('tema/contexto') or clean.lower().startswith('taxonomia') or clean.lower().startswith('observação'):
            bullets.append(clean)
    return bullets


def itemfmt_extract_block(lines, start_markers, stop_markers):
    start = -1
    for marker in start_markers:
        idx = itemfmt_find_line_idx(lines, marker)
        if idx != -1:
            start = idx
            break
    if start == -1:
        return []
    block = []
    for line in lines[start + 1:]:
        cl = itemfmt_clean_line(line)
        if any(cl.lower().startswith(m.lower()) for m in stop_markers):
            break
        block.append(line)
    return block


def itemfmt_parse_table_rows(lines):
    def is_header_triplet(parts):
        normalized = [itemfmt_clean_line(p).lower() for p in parts[:3]]
        return normalized == ['critério', 'status', 'evidência/ajuste']

    def is_markdown_separator(parts):
        if len(parts) < 3:
            return False
        return all(re.fullmatch(r':?-{3,}:?', itemfmt_clean_line(p)) for p in parts[:3])

    def parse_compact_single_line(text):
        text = itemfmt_clean_line(text)
        if not text:
            return []
        text = re.sub(r'\|---\|---\|---\|?', ' ', text)
        text = re.sub(r'Critério\s*Status\s*Evidência/Ajuste', '', text, flags=re.I)
        criterios = [
            'Unicidade do gabarito',
            'Clareza do comando',
            'Adequação ao ano-alvo',
            'Distratores plausíveis e distintos',
            'Aderência ao descritor/habilidade',
            'Restrições atendidas',
            '(Se houver imagem) Suporte necessário e reprodutível',
        ]
        criterio_pattern = '|'.join(re.escape(c) for c in criterios)
        status_pattern = r'(Atendido|Não se aplica|Parcialmente atendido|Não atendido)'
        pattern = re.compile(rf'({criterio_pattern})\s*({status_pattern})\s*(.*?)\s*(?=({criterio_pattern})\s*({status_pattern})|$)', re.I)
        found = []
        for m in pattern.finditer(text):
            criterio = itemfmt_clean_line(m.group(1))
            status = itemfmt_clean_line(m.group(2))
            evidencia = itemfmt_clean_line(m.group(3))
            if criterio and status and evidencia:
                found.append((criterio, status, evidencia))
        return found

    rows = []
    current = []
    for raw in lines:
        line = itemfmt_clean_line(raw)
        if not line:
            continue
        if line.lower() in {'critério', 'status', 'evidência/ajuste'}:
            continue

        stripped = raw.strip()
        if stripped.startswith('|') and '|' in stripped[1:]:
            parts = [itemfmt_clean_line(p) for p in stripped.strip('|').split('|')]
            parts = [p for p in parts if p]
            if len(parts) >= 3:
                if is_header_triplet(parts) or is_markdown_separator(parts):
                    continue
                rows.append((parts[0], parts[1], ' '.join(parts[2:])))
                continue

        if '	' in raw:
            parts = [itemfmt_clean_line(p) for p in raw.split('	') if itemfmt_clean_line(p)]
            if len(parts) >= 3:
                if is_header_triplet(parts) or is_markdown_separator(parts):
                    continue
                rows.append((parts[0], parts[1], ' '.join(parts[2:])))
                continue
        current.append(line)

    if rows:
        return [row for row in rows if not is_header_triplet(row) and not is_markdown_separator(row)]

    compact = [itemfmt_clean_line(x) for x in lines if itemfmt_clean_line(x)]
    compact = [x for x in compact if x.lower() not in {'critério', 'status', 'evidência/ajuste'}]
    compact = [x for x in compact if not re.fullmatch(r':?-{3,}:?', x)]

    if len(compact) == 1:
        single_line_rows = parse_compact_single_line(compact[0])
        if single_line_rows:
            return single_line_rows

    compact_joined = ' '.join(compact)
    single_line_rows = parse_compact_single_line(compact_joined)
    if single_line_rows:
        return single_line_rows

    i = 0
    while i + 2 < len(compact):
        triplet = (compact[i], compact[i + 1], compact[i + 2])
        if not is_header_triplet(triplet) and not is_markdown_separator(triplet):
            rows.append(triplet)
        i += 3
    return rows


def itemfmt_join_paragraph_lines(lines):
    parts = [itemfmt_clean_line(x) for x in lines if itemfmt_clean_line(x)]
    return ' '.join(parts)


def itemfmt_parse_lettered_blocks(lines):
    blocks = []
    current_letter = None
    current_text = []
    for raw in lines:
        line = itemfmt_clean_line(raw)
        if not line:
            continue
        m = re.match(r'^([A-D])\)\s*(.*)$', line)
        if m:
            if current_letter:
                blocks.append((current_letter, ' '.join(current_text).strip()))
            current_letter = m.group(1)
            current_text = [m.group(2).strip()] if m.group(2).strip() else []
        else:
            if current_letter:
                current_text.append(line)
    if current_letter:
        blocks.append((current_letter, ' '.join(current_text).strip()))
    return blocks


def itemfmt_parse_input_text(raw_text: str):
    lines = raw_text.replace('\r\n', '\n').replace('\r', '\n').split('\n')

    mini_relatorio_lines = itemfmt_extract_block(
        lines,
        ['MINI RELATÓRIO FINAL'],
        ['TABELA DE CHECAGEM TÉCNICA']
    )
    table_lines = itemfmt_extract_block(
        lines,
        ['TABELA DE CHECAGEM TÉCNICA'],
        ['ITEM FINAL']
    )
    item_lines = itemfmt_extract_block(
        lines,
        ['ITEM FINAL'],
        []
    )

    bullets = itemfmt_parse_bullets(mini_relatorio_lines)
    table_rows = itemfmt_parse_table_rows(table_lines)

    texto_base = itemfmt_join_paragraph_lines(itemfmt_extract_block(item_lines, ['Texto base'], ['Enunciado']))
    enunciado = itemfmt_join_paragraph_lines(itemfmt_extract_block(item_lines, ['Enunciado'], ['Comando']))
    comando = itemfmt_join_paragraph_lines(itemfmt_extract_block(item_lines, ['Comando'], ['Alternativas']))
    alternativas = itemfmt_parse_lettered_blocks(itemfmt_extract_block(item_lines, ['Alternativas'], ['Gabarito:']))

    gabarito_idx = itemfmt_find_line_idx(item_lines, 'Gabarito:')
    correct_letter = ''
    correct_just = ''
    if gabarito_idx != -1:
        gabarito_line = itemfmt_clean_line(item_lines[gabarito_idx])
        mg = re.search(r'Gabarito:\s*([A-D])', gabarito_line, flags=re.I)
        if mg:
            correct_letter = mg.group(1).upper()
        after = []
        for line in item_lines[gabarito_idx + 1:]:
            cl = itemfmt_clean_line(line)
            cl_lower = cl.lower()
            if cl_lower.startswith('justificativa técnica dos distratores'):
                break
            if cl_lower.startswith('justificativa do gabarito'):
                continue
            after.append(line)
        correct_just = itemfmt_join_paragraph_lines(after)

    distratores_lines = itemfmt_extract_block(item_lines, ['Justificativa técnica dos distratores'], ['Explicação ao aluno'])
    explicacao_lines = itemfmt_extract_block(item_lines, ['Explicação ao aluno'], [])

    return {
        'mini_relatorio': bullets,
        'tabela': table_rows,
        'texto_base': texto_base,
        'enunciado': enunciado,
        'comando': comando,
        'alternativas': alternativas,
        'gabarito': correct_letter,
        'justificativa_gabarito': correct_just,
        'justificativa_distratores': itemfmt_parse_lettered_blocks(distratores_lines),
        'explicacao_aluno': itemfmt_parse_lettered_blocks(explicacao_lines),
    }


def itemfmt_format_run(run, *, bold=False, underline=False, highlight=False, font_size=12):
    run.bold = bold
    run.underline = underline
    font = run.font
    font.name = 'Arial'
    font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), 'Arial')
    if highlight:
        rpr = run._r.get_or_add_rPr()
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), 'yellow')
        rpr.append(hl)


def itemfmt_add_normal_paragraph(doc, text='', space_after=4, first_line_cm=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if first_line_cm:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    r = p.add_run(text)
    itemfmt_format_run(r)
    return p


def itemfmt_add_heading(doc, text, level='main'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 'main' else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    if level == 'main':
        itemfmt_format_run(r, bold=True, font_size=14)
    else:
        itemfmt_format_run(r, bold=True, underline=True, font_size=12)
    return p


def itemfmt_set_table_autofit_window(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'autofit')

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), '5000')


def itemfmt_build_doc(data, output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(12)

    itemfmt_add_heading(doc, 'MINI RELATÓRIO FINAL', 'main')
    for item in data['mini_relatorio']:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        bullet = p.add_run('• ')
        itemfmt_format_run(bullet)
        txt = p.add_run(item)
        itemfmt_format_run(txt)

    itemfmt_add_heading(doc, 'TABELA DE CHECAGEM TÉCNICA', 'main')
    table = doc.add_table(rows=1, cols=3)
    itemfmt_set_table_autofit_window(table)
    widths = [Cm(8.0), Cm(2.7), Cm(11.3)]
    header = table.rows[0].cells
    titles = ['Critério', 'Status', 'Evidência/Ajuste']
    for i, cell in enumerate(header):
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(titles[i])
        itemfmt_format_run(r, bold=True)
        itemfmt_set_cell_border(cell)

    for idx, (criterio, status, evidencia) in enumerate(data['tabela']):
        row = table.add_row().cells
        values = [criterio, status, evidencia]
        for i, cell in enumerate(row):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(values[i])
            itemfmt_format_run(r, bold=(i == 0))
            itemfmt_set_cell_border(cell)
            if idx % 2 == 0:
                itemfmt_set_cell_shading(cell, 'D9D9D9')

    itemfmt_add_heading(doc, 'ITEM FINAL', 'main')
    if data.get('texto_base'):
        itemfmt_add_heading(doc, 'Texto base', 'sub')
        itemfmt_add_normal_paragraph(doc, data['texto_base'], space_after=8)
    itemfmt_add_heading(doc, 'Enunciado', 'sub')
    itemfmt_add_normal_paragraph(doc, data['enunciado'], space_after=8)
    itemfmt_add_heading(doc, 'Comando', 'sub')
    itemfmt_add_normal_paragraph(doc, data['comando'], space_after=8)
    itemfmt_add_heading(doc, 'Alternativas', 'sub')
    for letter, text in data['alternativas']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f'{letter}) ')
        itemfmt_format_run(r1, bold=True, highlight=(letter == data['gabarito']))
        r2 = p.add_run(text)
        itemfmt_format_run(r2, highlight=(letter == data['gabarito']))

    itemfmt_add_normal_paragraph(doc, f'Gabarito: {data["gabarito"]}', space_after=8)
    itemfmt_add_normal_paragraph(doc, data['justificativa_gabarito'], space_after=8)

    itemfmt_add_heading(doc, 'Justificativa técnica dos distratores', 'sub')
    for letter, text in data['justificativa_distratores']:
        itemfmt_add_normal_paragraph(doc, f'{letter}) {text}', space_after=8)

    itemfmt_add_heading(doc, 'Explicação ao aluno', 'sub')
    for letter, text in data['explicacao_aluno']:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r1 = p.add_run(f'{letter}) ')
        itemfmt_format_run(r1, bold=True)
        r2 = p.add_run(text)
        itemfmt_format_run(r2)

    doc.save(output_path)


class ItemFormatterTab(ttk.Frame):
    def __init__(self, master):
        super().__init__(master, padding=14)
        self.status_var = tk.StringVar(value='Pronto.')
        self._build_ui()

    def _build_ui(self):
        title = ttk.Label(self, text='Texto bruto → Word formatado', style="Title.TLabel")
        title.pack(anchor='w', pady=(0, 8))

        ttk.Label(
            self,
            text='Cole abaixo o texto no padrão MINI RELATÓRIO FINAL / TABELA DE CHECAGEM TÉCNICA / ITEM FINAL para gerar um .docx formatado.',
            style='Subtitle.TLabel',
            wraplength=1150,
        ).pack(anchor='w', pady=(0, 10))

        frame = ttk.Frame(self)
        frame.pack(fill='both', expand=True)

        topbar = ttk.Frame(frame)
        topbar.pack(fill='x', pady=(0, 8))

        ttk.Button(topbar, text='📂  Abrir TXT', command=self.load_txt, style='Soft.TButton').pack(side='left')
        ttk.Button(topbar, text='▶  Gerar Word formatado', command=self.generate_docx, style='Success.TButton').pack(side='left', padx=8)
        ttk.Button(topbar, text='🧹  Limpar', command=self.clear_text, style='Soft.TButton').pack(side='left')

        editor_wrap = ttk.Frame(frame)
        editor_wrap.pack(fill='both', expand=True)

        self.text = tk.Text(editor_wrap, wrap='word', font=('Segoe UI', 10), undo=True)
        self.text.pack(side='left', fill='both', expand=True)

        scroll = ttk.Scrollbar(editor_wrap, orient='vertical', command=self.text.yview)
        scroll.pack(side='right', fill='y')
        self.text.configure(yscrollcommand=scroll.set)

        footer = ttk.Frame(self)
        footer.pack(fill='x', pady=(8, 0))
        ttk.Label(footer, textvariable=self.status_var, style='Status.TLabel').pack(side='left')

    def load_txt(self):
        path = filedialog.askopenfilename(filetypes=[('Arquivos de texto', '*.txt'), ('Todos os arquivos', '*.*')])
        if not path:
            return
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        self.text.delete('1.0', 'end')
        self.text.insert('1.0', content)
        self.status_var.set(f'Texto carregado: {os.path.basename(path)}')

    def load_example(self):
        example = """MINI RELATÓRIO FINAL
· Matriz-base: SAEB 2001 – 5º ano – D20.
· Ano-alvo e dificuldade: 4º ano – Fácil.
· Tema/contexto e classificação do contexto (Pessoal/Ocupacional/Social/Científico): Exemplo – Social.
· Taxonomia Revisada: Aplicar + Procedimental.

TABELA DE CHECAGEM TÉCNICA
Critério
Status
Evidência/Ajuste
Unicidade do gabarito
Atendido
Há apenas uma alternativa compatível com os dados do item.
Clareza do comando
Atendido
O comando solicita diretamente a resposta esperada.

ITEM FINAL
Enunciado
Texto do enunciado aqui.
Comando
Texto do comando aqui.
Alternativas
A) Alternativa A.
B) Alternativa B.
C) Alternativa C.
D) Alternativa D.
Gabarito: C
Justificativa do gabarito aqui.
Justificativa técnica dos distratores
A) Justificativa técnica da alternativa A.
B) Justificativa técnica da alternativa B.
C) Justificativa técnica da alternativa C.
D) Justificativa técnica da alternativa D.
Explicação ao aluno
A) Explicação ao aluno da alternativa A.
B) Explicação ao aluno da alternativa B.
C) Explicação ao aluno da alternativa C.
D) Explicação ao aluno da alternativa D."""
        self.text.delete('1.0', 'end')
        self.text.insert('1.0', example)
        self.status_var.set('Modelo de exemplo inserido.')

    def clear_text(self):
        self.text.delete('1.0', 'end')
        self.status_var.set('Campo limpo.')

    def generate_docx(self):
        raw = self.text.get('1.0', 'end').strip()
        if not raw:
            messagebox.showwarning(ITEM_FORMATTER_TITLE, 'Cole um texto antes de gerar o Word.')
            return
        try:
            data = itemfmt_parse_input_text(raw)

            faltando = []
            if not data.get('enunciado'):
                faltando.append('Enunciado')
            if not data.get('comando'):
                faltando.append('Comando')

            alternativas = dict(data.get('alternativas', []))
            for letra in ['A', 'B', 'C', 'D']:
                if not alternativas.get(letra):
                    faltando.append(f'Alternativa {letra}')

            if not data.get('gabarito'):
                faltando.append('Gabarito')

            if faltando:
                raise ValueError(
                    'Não encontrei os campos essenciais do item. Faltando: ' + ', '.join(faltando)
                )

            save_path = filedialog.asksaveasfilename(
                defaultextension='.docx',
                initialfile=ITEM_FORMATTER_DEFAULT_FILENAME,
                filetypes=[('Word', '*.docx')]
            )
            if not save_path:
                return

            itemfmt_build_doc(data, save_path)
            self.status_var.set(f'Arquivo gerado: {save_path}')

            abrir = messagebox.askyesno(
                ITEM_FORMATTER_TITLE,
                f'Word gerado com sucesso:\n{save_path}\n\nDeseja abrir agora?'
            )
            if abrir:
                if os.name == 'nt':
                    os.startfile(save_path)
                elif sys.platform == 'darwin':
                    subprocess.run(['open', str(save_path)], check=False)
                else:
                    subprocess.run(['xdg-open', str(save_path)], check=False)
        except Exception as e:
            messagebox.showerror(ITEM_FORMATTER_TITLE, f'Erro ao gerar o arquivo:\n{e}')
            self.status_var.set('Falha ao gerar o arquivo.')




# ======================================================================
# BLOCO 4 - GERADOR DE TXT NO PADRÃO DE ITEM
# ======================================================================

TXT_GENERATOR_TEMPLATE ="""REGRAS INVIOLÁVEIS:
1 - Gerar o item internamente conforme os parâmetros fornecidos.
2 - Checar antes de enviar com a lista de verificação (descritor, ano-alvo, dificuldade, suporte, restrições, ordem crescente, gabarito pedido, distratores, formato).
3 - Se houver qualquer dúvida real (conflito de parâmetros, informação faltando, incompatibilidade entre gabarito e restrições, etc.), eu paro e pergunto objetivamente.
4 - Só depois de sanadas as dúvidas, eu envio a versão final.
5 - Instrução para copiar e colar:
	“Formatação: não usar quebras de linha manuais (^l). Usar apenas quebras de parágrafo (^p).”
	“Não inserir line breaks dentro de um mesmo parágrafo; cada nova linha deve ser um novo parágrafo.”

6 - A seção TABELA DE CHECAGEM TÉCNICA deve ser entregue exclusivamente com TAB real entre as colunas, dentro de um BLOCO DE CÓDIGO (code fence) usando três crases ``` na abertura e no fechamento.
	- É proibido usar | (pipes) em qualquer parte da tabela.
	- Deve haver exatamente 1 TAB entre cada coluna.
	- Não inserir quebras de linha dentro de uma mesma linha da tabela.
	- O conteúdo deve ser conciso; se necessário, resumir para caber na célula.
	- Não substituir TAB por espaços.

0) ITEM BASE (OBRIGATÓRIO — cole aqui o item completo)

Cole o item exatamente como está (inclua, se houver): suporte/descrição, Texto Base, enunciado, comando, alternativas A–D e gabarito.

ITEM BASE:
{item_base}

1) Parâmetros obrigatórios (preencher)
Componente: {componente}
Matriz-base: {matriz_base}
Etapa da matriz-base: {etapa}
Código-alvo: {codigo_alvo}
Ano-alvo do item: {ano_alvo}
Dificuldade: {dificuldade}
Tema/contexto: {tema}
Suporte: {suporte}
Restrições: {restricoes}
Gabarito desejado: {gabarito}

2) Regras de construção do item (obrigatórias)
Criar item inédito (novo do zero) a partir do ITEM BASE, sem copiar a mesma estrutura literal.
Avaliar UMA única habilidade/descritor: exatamente o código informado.
Manter a mesma intenção avaliativa do código-alvo, mas com novo contexto e/ou novos números (respeitando restrições).
Linguagem adequada ao ano-alvo, formal, clara; sem "pegadinhas"; sem ambiguidade.
Gabarito único: revisar para garantir que não há outra alternativa correta.
Alternativas A–D: 1 correta + 3 distratores plausíveis, cada um representando um erro típico diferente.
Ordem crescente nas alternativas (quando forem numéricas).
Se a alternativa correta ficar muito distante das demais, ajuste valores para aproximar, mantendo plausibilidade dos erros.
Respeitar integralmente as restrições (intervalos, ordens, unidades, proibições).
Se ano-alvo ≠ etapa da matriz-base, aplicar adaptação controlada: ajustar linguagem, números e etapas de resolução, sem mudar o foco do código.

3) Suporte (imagem) — MODELO CONDICIONAL

Regra: só usar esta seção se Suporte: com imagem.
Se Suporte: sem imagem, não gerar descrição/prompt de imagem.

PROMPT PADRÃO PARA CRIAÇÃO DE IMAGENS (SUPORTE DE ITEM)
Objetivo: criar um suporte em imagem necessário para resolver o item, sem ambiguidade e sem dar pistas do gabarito.

A) Metadados do item

Ano-alvo do estudante: [{ano_alvo}]
Componente: [{componente}]
Matriz-base/código: [{matriz_base} / {etapa} / {codigo_alvo}]
Tema/contexto: [{tema}]
Restrições numéricas/textuais: [{restricoes}]

B) Briefing visual

Cenário: [{{onde acontece}}]
Objetos essenciais: [{{lista}}]
Objetos proibidos: calculadora; "total"; "resultado"; "resposta"; "correto"; setas/destaques; pistas do gabarito.

C) Conteúdo informacional obrigatório (preciso e reprodutível)

Quantidade exata de elementos: [{{n}}]
Ordem/posição (esquerda→direita / cima→baixo): [{{mapa posição→rótulo}}]
Rótulos/valores (exatamente assim): [{{lista}}]
Formatação: fonte legível; alto contraste; sem números extras.

D) Antiambiguidade / Anti-gabarito
Sem marcações; sem inferência; sem destaque do correto; escala explícita se houver.

E) Checklist
Imagem contém somente o necessário; tudo legível; sem pistas; dados suficientes.

4) Formato de entrega (OBRIGATÓRIO) — responder exatamente assim

MINI RELATÓRIO FINAL
• Matriz-base: (matriz + etapa + código + descrição)
• Ano-alvo e dificuldade
• Tema/contexto e classificação do contexto (Pessoal/Ocupacional/Social/Científico)
• Taxonomia Revisada: processo cognitivo + dimensão do conhecimento
• Observação de adaptação (se ano-alvo ≠ etapa)

TABELA DE CHECAGEM TÉCNICA
```text
Critério	Status	Evidência/Ajuste
Unicidade do gabarito		
Clareza do comando		
Adequação ao ano-alvo		
Distratores plausíveis e distintos		
Aderência ao descritor/habilidade		
Restrições atendidas		
(Se houver imagem) Suporte necessário e reprodutível		
````

ITEM FINAL

Suporte (somente se houver; se "sem imagem", omitir)
Texto Base (se houver; senão omitir)
Enunciado
Comando
Alternativas
Gabarito: X

Justificativa do gabarito (OBRIGATÓRIA)

Começar exatamente com: "Está correta. O aluno provavelmente "
Incluir, obrigatoriamente:

Leitura matemática do contexto (ex.: "repartir igualmente", "grupos iguais", "configuração retangular", "comparação multiplicativa", "combinatória").
Operação escolhida e por quê (multiplicação ou divisão) em 1 frase.
Cálculo essencial (sem excesso de passos, coerente com a dificuldade).
Checagem/validação do resultado (operação inversa: quociente × divisor = total; ou estimativa).
Por que é único (não há outra alternativa que satisfaça simultaneamente os dados do enunciado).
Vínculo explícito com o descritor/código (mencionar Dxx/Hxx).

Justificativa técnica dos distratores (OBRIGATÓRIA)

Para cada alternativa incorreta, começar exatamente com: "Está incorreta. O aluno provavelmente "
E fundamentar com:

Erro típico principal (distinto):
erro de interpretação do contexto (ex.: confundiu "grupos" com "total"), ou
escolha inadequada da operação (multiplicou quando deveria dividir / dividiu por outro número), ou
erro procedimental de cálculo (tabuada, algoritmo da divisão), ou
uso parcial/indevido de um dado (ignorou um valor, trocou divisor).
Mostre a consequência numérica em 1 linha (ex.: "fez 360÷10", "fez 12×30", "usou metade/dobro").
Teste rápido de inconsistência (operação inversa ou comparação com o total: "se fosse X, então X×12 daria …, que não é …").

EXPLICAÇÃO AO ALUNO (OBRIGATÓRIA)

A) Está incorreta. Se você marcou esta alternativa possivelmente …
[explicar em linguagem didática: o que o problema pede; por que a operação faz sentido; uma forma rápida de conferir. Evitar "a resposta é…". Limite: ~450 caracteres.]

B) Está incorreta. Se você marcou esta alternativa possivelmente …
[idem]

C) Está incorreta. Se você marcou esta alternativa possivelmente …
[idem]

D) Está incorreta. Se você marcou esta alternativa possivelmente …
[idem]

Observações de qualidade (para orientar a geração)
Não copiar a estrutura literal do item-base; mudar cenário e/ou organização dos dados.
Um item = um descritor (não misturar habilidades).
Alternativas numéricas em ordem crescente.
Distratores próximos ao gabarito quando possível (para aumentar plausibilidade), mas sempre com erros típicos diferentes.
Respeitar limites de extensão e restrições numéricas/textuais.
"""


class TxtGeneratorTab(ttk.Frame):
    def __init__(self, master):
        super().__init__(master, padding=14)
        self._build_ui()

    def _build_ui(self):
        ttk.Label(self, text="Gerador de TXT — Padrão de item", style="Title.TLabel").pack(anchor="w", pady=(0, 10))
        ttk.Label(self, text="Preencha os campos, cole o item base e gere o texto no padrão solicitado.", style="Subtitle.TLabel", wraplength=1200).pack(anchor="w", pady=(0, 10))

        main = ttk.Frame(self)
        main.pack(fill="both", expand=True)

        left = ttk.Frame(main)
        left.pack(side="left", fill="both", expand=False, padx=(0, 10))
        right = ttk.Frame(main)
        right.pack(side="left", fill="both", expand=True)

        form = ttk.Frame(left)
        form.pack(fill="x")

        self.componente = tk.StringVar(value="Matemática")
        self.matriz_base = tk.StringVar(value="SAEB 2001")
        self.etapa = tk.StringVar(value="5º ano")
        self.codigo_alvo = tk.StringVar()
        self.ano_alvo = tk.StringVar(value="4º ano")
        self.dificuldade = tk.StringVar(value="Média")
        self.tema = tk.StringVar()
        self.suporte = tk.StringVar(value="sem imagem")
        self.restricoes = tk.StringVar()
        self.gabarito = tk.StringVar(value="A")
        self.status_var = tk.StringVar(value="Pronto.")

        self._add_combo(form, "Componente", self.componente, ["Matemática", "Língua Portuguesa"], 0)
        self._add_combo(form, "Matriz-base", self.matriz_base, ["SAEB 2001", "SAEB 2018", "SPAECE"], 1)
        self._add_combo(form, "Etapa da matriz-base", self.etapa, ["2º ano", "5º ano", "9º ano"], 2)
        self._add_entry(form, "Código-alvo", self.codigo_alvo, 3)
        self._add_entry(form, "Ano-alvo do item", self.ano_alvo, 4)
        self._add_combo(form, "Dificuldade", self.dificuldade, ["Fácil", "Média", "Alta", "Difícil"], 5)
        self._add_entry(form, "Tema/contexto", self.tema, 6)
        self._add_combo(form, "Suporte", self.suporte, ["sem imagem", "com imagem", "tabela", "gráfico", "malha"], 7)
        self._add_entry(form, "Restrições", self.restricoes, 8)
        self._add_combo(form, "Gabarito desejado", self.gabarito, ["A", "B", "C", "D"], 9)

        ttk.Label(left, text="ITEM BASE (cole aqui o item completo)", font=("Segoe UI Semibold", 11, "bold")).pack(anchor="w", pady=(14, 6))
        self.item_text = tk.Text(left, wrap="word", height=20, font=("Segoe UI", 10))
        self.item_text.pack(fill="both", expand=True)

        btns = ttk.Frame(left)
        btns.pack(fill="x", pady=(10, 0))
        ttk.Button(btns, text="▶  Gerar TXT", command=self.gerar_txt, style="Success.TButton").pack(side="left", padx=(0, 6))
        ttk.Button(btns, text="🧹  Limpar", command=self.limpar, style="Soft.TButton").pack(side="left", padx=(0, 6))
        ttk.Button(btns, text="📋  Copiar saída", command=self.copiar_saida, style="Soft.TButton").pack(side="left", padx=(0, 6))
        ttk.Button(btns, text="💾  Salvar TXT", command=self.salvar_txt, style="Accent.TButton").pack(side="left")

        ttk.Label(right, text="TXT gerado", font=("Segoe UI Semibold", 11, "bold")).pack(anchor="w", pady=(0, 8))
        self.output_text = tk.Text(right, wrap="word", font=("Consolas", 10), relief="flat", bd=0, highlightthickness=1)
        self.output_text.pack(fill="both", expand=True)

        footer = ttk.Frame(self)
        footer.pack(fill='x', pady=(8, 0))
        ttk.Label(footer, textvariable=self.status_var, style='Status.TLabel').pack(side='left')

    def _add_entry(self, parent, label, variable, row):
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", padx=(0, 10), pady=4)
        entry = ttk.Entry(parent, textvariable=variable, width=48)
        entry.grid(row=row, column=1, sticky="ew", pady=4)
        parent.columnconfigure(1, weight=1)

    def _add_combo(self, parent, label, variable, values, row):
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", padx=(0, 10), pady=4)
        combo = ttk.Combobox(parent, textvariable=variable, values=values, state="readonly", width=45)
        combo.grid(row=row, column=1, sticky="ew", pady=4)
        parent.columnconfigure(1, weight=1)

    def _collect_data(self):
        return {
            "item_base": self.item_text.get("1.0", "end").strip(),
            "componente": self.componente.get().strip(),
            "matriz_base": self.matriz_base.get().strip(),
            "etapa": self.etapa.get().strip(),
            "codigo_alvo": self.codigo_alvo.get().strip(),
            "ano_alvo": self.ano_alvo.get().strip(),
            "dificuldade": self.dificuldade.get().strip(),
            "tema": self.tema.get().strip(),
            "suporte": self.suporte.get().strip(),
            "restricoes": self.restricoes.get().strip(),
            "gabarito": self.gabarito.get().strip().upper(),
        }

    def gerar_txt(self):
        data = self._collect_data()
        if not data["item_base"]:
            messagebox.showwarning("Aviso", "Cole o item completo no campo ITEM BASE.")
            return
        if not data["codigo_alvo"]:
            messagebox.showwarning("Aviso", "Preencha o campo Código-alvo.")
            return
        generated = TXT_GENERATOR_TEMPLATE.format(**data)
        self.output_text.delete("1.0", "end")
        self.output_text.insert("1.0", generated)
        self.status_var.set("TXT gerado.")

    def limpar(self):
        defaults = [
            (self.componente, "Matemática"),
            (self.matriz_base, "SAEB 2001"),
            (self.etapa, "5º ano"),
            (self.codigo_alvo, ""),
            (self.ano_alvo, "4º ano"),
            (self.dificuldade, "Média"),
            (self.tema, ""),
            (self.suporte, "sem imagem"),
            (self.restricoes, ""),
            (self.gabarito, "A"),
        ]
        for var, default in defaults:
            var.set(default)
        self.item_text.delete("1.0", "end")
        self.output_text.delete("1.0", "end")
        self.status_var.set("Campos limpos.")

    def copiar_saida(self):
        text = self.output_text.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("Aviso", "Gere o TXT antes de copiar.")
            return
        self.clipboard_clear()
        self.clipboard_append(text)
        self.update()
        self.status_var.set("Conteúdo copiado para a área de transferência.")
        messagebox.showinfo("Copiado", "O conteúdo gerado foi copiado para a área de transferência.")

    def salvar_txt(self):
        text = self.output_text.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("Aviso", "Gere o TXT antes de salvar.")
            return
        nome_padrao = f"item_padrao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            initialfile=nome_padrao,
            filetypes=[("Arquivo de texto", "*.txt"), ("Todos os arquivos", "*.*")],
        )
        if not file_path:
            return
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(text)
        self.status_var.set(f"Arquivo salvo: {file_path}")
        messagebox.showinfo("Salvo", f"Arquivo salvo em:\n{file_path}")


# ======================================================================
# TEMA DA INTERFACE
# ======================================================================

THEMES = {
    "light": {
        "bg":                "#F5F7FB",
        "fg":                "#1C2635",
        "card":              "#FFFFFF",
        "card_border":       "#DDE4F0",
        "input_bg":          "#FFFFFF",
        "input_border":      "#C4CFDF",
        "muted":             "#64748B",
        "select_bg":         "#BFDBFE",
        "select_fg":         "#0C1A3A",
        "tab_active_bg":     "#1D4ED8",
        "tab_active_fg":     "#FFFFFF",
        "tab_inactive_bg":   "#E8EEF9",
        "tab_inactive_fg":   "#334155",
        "tab_hover_bg":      "#D0DBEF",
        "btn_primary_bg":    "#2563EB",
        "btn_primary_fg":    "#FFFFFF",
        "btn_primary_hover": "#1D4ED8",
        "btn_secondary_bg":  "#EDF0F7",
        "btn_secondary_fg":  "#1C2635",
        "btn_secondary_hover": "#D8DEF0",
        "btn_danger_bg":     "#EF4444",
        "btn_danger_fg":     "#FFFFFF",
        "btn_danger_hover":  "#DC2626",
        "btn_success_bg":    "#16A34A",
        "btn_success_fg":    "#FFFFFF",
        "btn_success_hover": "#15803D",
        "progress_trough":   "#DDE4F0",
        "progress_fill":     "#2563EB",
        "drop_bg":           "#EEF4FF",
        "drop_border":       "#BFCFEE",
        "status_fg":         "#64748B",
        "scrollbar_bg":      "#DDE4F0",
        "topbar_bg":         "#1C2635",
        "topbar_fg":         "#FFFFFF",
        "separator":         "#E2E8F0",
        "badge_bg":          "#DBEAFE",
        "badge_fg":          "#1D4ED8",
    },
    "dark": {
        "bg":                "#0D1525",
        "fg":                "#E2E8F4",
        "card":              "#162036",
        "card_border":       "#263D65",
        "input_bg":          "#0F1C34",
        "input_border":      "#263D65",
        "muted":             "#8899BB",
        "select_bg":         "#1E3A8A",
        "select_fg":         "#E2E8F4",
        "tab_active_bg":     "#2563EB",
        "tab_active_fg":     "#FFFFFF",
        "tab_inactive_bg":   "#1A2B4A",
        "tab_inactive_fg":   "#B8C8E0",
        "tab_hover_bg":      "#22375E",
        "btn_primary_bg":    "#2563EB",
        "btn_primary_fg":    "#FFFFFF",
        "btn_primary_hover": "#3B82F6",
        "btn_secondary_bg":  "#1A2B4A",
        "btn_secondary_fg":  "#C5D3E8",
        "btn_secondary_hover": "#22375E",
        "btn_danger_bg":     "#DC2626",
        "btn_danger_fg":     "#FFFFFF",
        "btn_danger_hover":  "#B91C1C",
        "btn_success_bg":    "#15803D",
        "btn_success_fg":    "#FFFFFF",
        "btn_success_hover": "#166534",
        "progress_trough":   "#1A2B4A",
        "progress_fill":     "#3B82F6",
        "drop_bg":           "#111E35",
        "drop_border":       "#263D65",
        "status_fg":         "#8899BB",
        "scrollbar_bg":      "#1A2B4A",
        "topbar_bg":         "#0A1020",
        "topbar_fg":         "#E2E8F4",
        "separator":         "#1E2D48",
        "badge_bg":          "#1E3A8A",
        "badge_fg":          "#93C5FD",
    },
}

# Referência global do tema atual (acessível por todas as abas)
_CURRENT_THEME: dict = {}


def apply_app_theme(root, mode="light"):
    global _CURRENT_THEME
    colors = THEMES.get(mode, THEMES["light"])
    _CURRENT_THEME = colors

    style = ttk.Style(root)
    try:
        if "clam" in style.theme_names():
            style.theme_use("clam")
    except Exception:
        pass

    try:
        root.configure(bg=colors["bg"])
    except Exception:
        pass

    # --- Base ---
    style.configure(".",
        background=colors["bg"],
        foreground=colors["fg"],
        font=("Segoe UI", 10),
        bordercolor=colors["card_border"],
    )

    # --- Frames ---
    style.configure("TFrame", background=colors["bg"])
    style.configure("CardInner.TFrame", background=colors["card"])

    # --- Labels ---
    style.configure("TLabel",
        background=colors["bg"],
        foreground=colors["fg"],
        font=("Segoe UI", 10),
    )
    style.configure("Title.TLabel",
        background=colors["bg"],
        foreground=colors["fg"],
        font=("Segoe UI Semibold", 15, "bold"),
    )
    style.configure("Subtitle.TLabel",
        background=colors["bg"],
        foreground=colors["muted"],
        font=("Segoe UI", 10),
    )
    style.configure("Status.TLabel",
        background=colors["bg"],
        foreground=colors["status_fg"],
        font=("Segoe UI", 10),
    )
    style.configure("Mini.TLabel",
        background=colors["card"],
        foreground=colors["muted"],
        font=("Segoe UI", 9),
    )
    style.configure("MiniOnBg.TLabel",
        background=colors["bg"],
        foreground=colors["muted"],
        font=("Segoe UI", 9),
    )

    # --- LabelFrame (cards) ---
    style.configure("TLabelframe",
        background=colors["card"],
        foreground=colors["fg"],
        bordercolor=colors["card_border"],
        relief="solid",
        borderwidth=1,
    )
    style.configure("TLabelframe.Label",
        background=colors["card"],
        foreground=colors["fg"],
        font=("Segoe UI Semibold", 10, "bold"),
    )
    style.configure("Card.TLabelframe",
        background=colors["card"],
        foreground=colors["fg"],
        bordercolor=colors["card_border"],
        relief="solid",
        borderwidth=1,
    )
    style.configure("Card.TLabelframe.Label",
        background=colors["card"],
        foreground=colors["fg"],
        font=("Segoe UI Semibold", 10, "bold"),
    )

    # --- Botões ---
    _btn_pad = (14, 7)

    style.configure("TButton",
        background=colors["btn_secondary_bg"],
        foreground=colors["btn_secondary_fg"],
        font=("Segoe UI", 10),
        padding=_btn_pad,
        relief="flat",
        borderwidth=0,
    )
    style.map("TButton",
        background=[
            ("active", colors["btn_secondary_hover"]),
            ("pressed", colors["btn_secondary_hover"]),
            ("disabled", colors["card_border"]),
        ],
        foreground=[
            ("disabled", colors["muted"]),
        ],
    )

    style.configure("Accent.TButton",
        background=colors["btn_primary_bg"],
        foreground=colors["btn_primary_fg"],
        font=("Segoe UI Semibold", 10),
        padding=_btn_pad,
        relief="flat",
        borderwidth=0,
    )
    style.map("Accent.TButton",
        background=[
            ("active", colors["btn_primary_hover"]),
            ("pressed", colors["btn_primary_hover"]),
            ("disabled", colors["card_border"]),
        ],
        foreground=[
            ("active", colors["btn_primary_fg"]),
            ("pressed", colors["btn_primary_fg"]),
            ("disabled", colors["muted"]),
            ("!disabled", colors["btn_primary_fg"]),
        ],
    )

    style.configure("Soft.TButton",
        background=colors["btn_secondary_bg"],
        foreground=colors["btn_secondary_fg"],
        font=("Segoe UI", 10),
        padding=_btn_pad,
        relief="flat",
        borderwidth=0,
    )
    style.map("Soft.TButton",
        background=[
            ("active", colors["btn_secondary_hover"]),
            ("pressed", colors["btn_secondary_hover"]),
            ("disabled", colors["card_border"]),
        ],
        foreground=[
            ("active", colors["btn_secondary_fg"]),
            ("pressed", colors["btn_secondary_fg"]),
            ("disabled", colors["muted"]),
            ("!disabled", colors["btn_secondary_fg"]),
        ],
    )

    style.configure("Danger.TButton",
        background=colors["btn_danger_bg"],
        foreground=colors["btn_danger_fg"],
        font=("Segoe UI", 10),
        padding=_btn_pad,
        relief="flat",
        borderwidth=0,
    )
    style.map("Danger.TButton",
        background=[
            ("active", colors["btn_danger_hover"]),
            ("pressed", colors["btn_danger_hover"]),
            ("disabled", colors["card_border"]),
        ],
        foreground=[
            ("active", colors["btn_danger_fg"]),
            ("disabled", colors["muted"]),
            ("!disabled", colors["btn_danger_fg"]),
        ],
    )

    style.configure("Success.TButton",
        background=colors["btn_success_bg"],
        foreground=colors["btn_success_fg"],
        font=("Segoe UI Semibold", 10),
        padding=_btn_pad,
        relief="flat",
        borderwidth=0,
    )
    style.map("Success.TButton",
        background=[
            ("active", colors["btn_success_hover"]),
            ("pressed", colors["btn_success_hover"]),
            ("disabled", colors["card_border"]),
        ],
        foreground=[
            ("active", colors["btn_success_fg"]),
            ("disabled", colors["muted"]),
            ("!disabled", colors["btn_success_fg"]),
        ],
    )

    # --- Inputs ---
    style.configure("TEntry",
        fieldbackground=colors["input_bg"],
        foreground=colors["fg"],
        bordercolor=colors["input_border"],
        insertcolor=colors["fg"],
    )
    style.configure("TSpinbox",
        fieldbackground=colors["input_bg"],
        foreground=colors["fg"],
        bordercolor=colors["input_border"],
        arrowcolor=colors["fg"],
    )
    style.configure("TCombobox",
        fieldbackground=colors["input_bg"],
        foreground=colors["fg"],
        bordercolor=colors["input_border"],
        arrowcolor=colors["fg"],
        selectbackground=colors["select_bg"],
        selectforeground=colors["select_fg"],
    )
    style.map("TCombobox",
        fieldbackground=[("readonly", colors["input_bg"])],
        foreground=[("readonly", colors["fg"])],
    )

    # --- Checkbutton / Radiobutton ---
    style.configure("TCheckbutton",
        background=colors["bg"],
        foreground=colors["fg"],
        font=("Segoe UI", 10),
    )
    style.configure("TRadiobutton",
        background=colors["bg"],
        foreground=colors["fg"],
        font=("Segoe UI", 10),
    )

    # --- Notebook (abas) ---
    style.configure("TNotebook",
        background=colors["bg"],
        borderwidth=0,
        tabmargins=(0, 0, 0, 0),
    )
    style.configure("TNotebook.Tab",
        background=colors["tab_inactive_bg"],
        foreground=colors["tab_inactive_fg"],
        font=("Segoe UI", 10),
        padding=(18, 9),
    )
    style.map("TNotebook.Tab",
        background=[
            ("selected", colors["tab_active_bg"]),
            ("active",   colors["tab_hover_bg"]),
        ],
        foreground=[
            ("selected", colors["tab_active_fg"]),
            ("active",   colors["tab_inactive_fg"]),
        ],
    )

    # --- Progressbar ---
    style.configure("Horizontal.TProgressbar",
        thickness=12,
        troughcolor=colors["progress_trough"],
        background=colors["progress_fill"],
        bordercolor=colors["progress_trough"],
        lightcolor=colors["progress_fill"],
        darkcolor=colors["progress_fill"],
    )

    # --- Scrollbar ---
    style.configure("TScrollbar",
        background=colors["scrollbar_bg"],
        troughcolor=colors["bg"],
        bordercolor=colors["bg"],
        arrowcolor=colors["muted"],
    )

    # --- Aplicar a widgets tk nativos (Text, Listbox, Frame, etc.) ---
    def _apply_to_widget(widget):
        try:
            cls = widget.winfo_class()
        except Exception:
            cls = ""
        try:
            if cls in {"Text", "Listbox"}:
                widget.configure(
                    bg=colors["input_bg"],
                    fg=colors["fg"],
                    insertbackground=colors["fg"],
                    selectbackground=colors["select_bg"],
                    selectforeground=colors["select_fg"],
                    highlightbackground=colors["input_border"],
                    highlightcolor=colors["btn_primary_bg"],
                )
            elif cls == "Canvas":
                widget.configure(
                    bg=colors["bg"],
                    highlightbackground=colors["card_border"],
                )
            elif cls in {"Frame", "LabelFrame"}:
                widget.configure(bg=colors["bg"])
            elif cls == "Label":
                widget.configure(bg=colors["bg"], fg=colors["fg"])
            elif cls in {"Entry", "Spinbox"}:
                widget.configure(
                    bg=colors["input_bg"],
                    fg=colors["fg"],
                    insertbackground=colors["fg"],
                    highlightbackground=colors["input_border"],
                    highlightcolor=colors["btn_primary_bg"],
                )
        except Exception:
            pass
        for child in widget.winfo_children():
            _apply_to_widget(child)

    _apply_to_widget(root)


class CombinedApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Ferramentas SAFA")
        self.root.geometry("1400x900")
        self.root.minsize(1100, 760)
        self.theme_var = tk.StringVar(value="Claro")
        try:
            self.root.state("zoomed")
        except Exception:
            try:
                self.root.attributes("-zoomed", True)
            except Exception:
                pass

        self._build_ui()
        self._on_theme_change()

    def _build_ui(self):
        # ── Topbar fixa (fundo escuro) ─────────────────────────────────
        self._topbar_frame = tk.Frame(self.root, height=52)
        self._topbar_frame.pack(fill="x", side="top")
        self._topbar_frame.pack_propagate(False)

        # Logo / título
        self._title_lbl = tk.Label(
            self._topbar_frame,
            text="⚙  Ferramentas SAFA",
            font=("Segoe UI Semibold", 13, "bold"),
            padx=18,
        )
        self._title_lbl.pack(side="left", pady=10)

        # Versão / badge
        self._badge_lbl = tk.Label(
            self._topbar_frame,
            text=" v2 ",
            font=("Segoe UI", 8, "bold"),
            padx=6, pady=2,
            relief="flat",
        )
        self._badge_lbl.pack(side="left", pady=16)

        # Tema (direita)
        theme_frame = tk.Frame(self._topbar_frame)
        theme_frame.pack(side="right", padx=18)

        tk.Label(
            theme_frame,
            text="🌓",
            font=("Segoe UI", 11),
        ).pack(side="left", padx=(0, 4))

        self._theme_combo = ttk.Combobox(
            theme_frame,
            textvariable=self.theme_var,
            values=["Claro", "Escuro"],
            state="readonly",
            width=8,
            font=("Segoe UI", 10),
        )
        self._theme_combo.pack(side="left")
        self._theme_combo.bind("<<ComboboxSelected>>", lambda e: self._on_theme_change())

        # ── Conteúdo (notebook) ────────────────────────────────────────
        container = ttk.Frame(self.root, padding=(10, 8, 10, 6))
        container.pack(fill="both", expand=True)

        notebook = ttk.Notebook(container)
        notebook.pack(fill="both", expand=True)

        self.txt_generator_tab = TxtGeneratorTab(notebook)
        self.item_formatter_tab = ItemFormatterTab(notebook)
        self.safa_tab = SafaTab(notebook)
        self.word_png_tab = WordPngTab(notebook)

        notebook.add(self.txt_generator_tab,   text="  📝  Gerador de TXT  ")
        notebook.add(self.item_formatter_tab,  text="  🗂  Formatador de item  ")
        notebook.add(self.safa_tab,            text="  🔄  Padrão SAFA  ")
        notebook.add(self.word_png_tab,        text="  🖼  Montagem das imagens  ")

    def _on_theme_change(self):
        mode = "dark" if self.theme_var.get() == "Escuro" else "light"
        apply_app_theme(self.root, mode)
        c = _CURRENT_THEME
        # Atualiza topbar manualmente (tk.Frame, não ttk)
        try:
            self._topbar_frame.configure(bg=c["topbar_bg"])
            self._title_lbl.configure(bg=c["topbar_bg"], fg=c["topbar_fg"])
            self._badge_lbl.configure(bg=c["badge_bg"], fg=c["badge_fg"])
            for child in self._topbar_frame.winfo_children():
                if isinstance(child, tk.Frame):
                    child.configure(bg=c["topbar_bg"])
                    for sub in child.winfo_children():
                        if isinstance(sub, tk.Label):
                            sub.configure(bg=c["topbar_bg"], fg=c["topbar_fg"])
        except Exception:
            pass


def main():
    root_class = TkinterDnD.Tk if DND_AVAILABLE else tk.Tk
    root = root_class()
    CombinedApp(root)
    root.mainloop()


if __name__ == "__main__":
    try:
        main()
    except Exception:
        msg = traceback.format_exc()
        show_fatal_error_and_wait(
            "O programa encontrou um erro inesperado ao iniciar:\n\n" + msg
        )




