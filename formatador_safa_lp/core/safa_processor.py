import re
import json
from pathlib import Path
from docx import Document
from datetime import datetime

class SafaProcessor:
    def __init__(self):
        self.base_path = Path("formatador_safa_lp")
        self.logs_dir = self.base_path / "logs"

    def parse_intervalo_itens(self, entrada_texto: str) -> set[int]:
        if not entrada_texto or not entrada_texto.strip():
            return set()
        
        ids = set()
        # Divide por vírgula ou ponto e vírgula
        partes = re.split(r'[,;]', entrada_texto)
        
        for parte in partes:
            parte = parte.strip()
            if '-' in parte:
                try:
                    inicio, fim = map(int, parte.split('-'))
                    ids.update(range(inicio, fim + 1))
                except ValueError:
                    continue
            else:
                try:
                    ids.add(int(parte))
                except ValueError:
                    continue
        return ids

    def processar_lote(self, itens_parseados: list[dict], itens_filtrados: set[int], configs: dict) -> dict:
        pasta_saida = Path(configs.get('pasta_saida', self.base_path / "output"))
        if not pasta_saida.exists():
            pasta_saida.mkdir(parents=True, exist_ok=True)
            
        relatorio = {"data": datetime.now().isoformat(), "itens": []}
        campos_obrigatorios = ["enunciado", "comando", "alternativa_a", "alternativa_b", "alternativa_c", "alternativa_d", "gabarito", "explicacao_aluno"]
        
        for item in itens_parseados:
            numero = item.get('numero')
            if itens_filtrados and numero not in itens_filtrados:
                continue
                
            status_item = {"numero": numero, "status": "sucesso", "erros": []}
            
            # Validação
            for campo in campos_obrigatorios:
                if not item.get(campo):
                    status_item["erros"].append(f"Campo obrigatório ausente: {campo}")
            
            if item.get("texto_base") and not item.get("fonte_autor"):
                status_item["erros"].append("Fonte/Autor obrigatório para itens com Texto-base.")
                
            if status_item["erros"]:
                status_item["status"] = "erro"
                relatorio["itens"].append(status_item)
                continue
                
            try:
                caminho_arquivo = pasta_saida / f"item_{numero:02d}_padronizado.docx"
                self.gerar_documento_6_paginas(item, configs.get('incluir_ensinart', False), str(caminho_arquivo))
                relatorio["itens"].append(status_item)
            except Exception as e:
                status_item["status"] = "erro"
                status_item["erros"].append(str(e))
                relatorio["itens"].append(status_item)
                
        # Salva relatório
        if not self.logs_dir.exists():
            self.logs_dir.mkdir(parents=True, exist_ok=True)
        with open(self.logs_dir / "relatorio_processamento.json", "w", encoding="utf-8") as f:
            json.dump(relatorio, f, indent=4, ensure_ascii=False)
            
        return relatorio

    def gerar_documento_6_paginas(self, item: dict, incluir_ensinart: bool, caminho_saida: str):
        doc = Document()
        
        # Página 1: Capa/Enunciado
        if incluir_ensinart:
            doc.add_paragraph("ENSINART")
        
        if item.get("texto_base"):
            doc.add_paragraph(item["texto_base"])
            if item.get("fonte_autor"):
                p_fonte = doc.add_paragraph(item["fonte_autor"])
                p_fonte.alignment = 2 # Direita
                
        if item.get("suporte"):
            doc.add_paragraph(f"Suporte: {item['suporte']}")
            
        doc.add_paragraph(item["enunciado"])
        doc.add_paragraph(item["comando"])
        doc.add_page_break()
        
        # Páginas 2-5: Alternativas
        for letra in ["a", "b", "c", "d"]:
            texto_alt = item[f"alternativa_{letra}"]
            # Remove prefixo se existir (ex: "A) ")
            texto_limpo = re.sub(rf"^{letra.upper()}\)\s*", "", texto_alt, flags=re.IGNORECASE)
            doc.add_paragraph(texto_limpo)
            doc.add_page_break()
            
        # Página 6: Explicação
        doc.add_paragraph("EXPLICAÇÃO UNIFICADA")
        gabarito = item["gabarito"].strip().upper()
        
        for letra in ["A", "B", "C", "D"]:
            texto_alt = item[f"alternativa_{letra.lower()}"]
            texto_limpo = re.sub(rf"^{letra}\)\s*", "", texto_alt, flags=re.IGNORECASE)
            status = "CORRETA" if letra == gabarito else "INCORRETA"
            # Aqui assume-se que a explicação no item_parser já vem formatada ou é genérica
            expl = item.get("explicacao_aluno", "Sem explicação detalhada.")
            doc.add_paragraph(f'"{texto_limpo}" ESTÁ {status}. {expl}')
            
        doc.save(caminho_saida)
