import os
import sys
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt

class WordFormatter:
    def __init__(self):
        self.base_path = Path("formatador_safa_lp")
        self.output_dir = self.base_path / "output"
        self.sections_obrigatorias = [
            "MINI RELATÓRIO FINAL", "TABELA DE CHECAGEM TÉCNICA", "ITEM FINAL", 
            "Enunciado", "Comando", "Alternativas", "A)", "B)", "C)", "D)", 
            "Gabarito", "JUSTIFICATIVA DO GABARITO", 
            "JUSTIFICATIVA TÉCNICA DOS DISTRATORES", "EXPLICAÇÃO AO ALUNO"
        ]

    def validar_estrutura(self, texto_bruto: str) -> tuple[bool, str]:
        for secao in self.sections_obrigatorias:
            if secao not in texto_bruto:
                return False, f"Não foi possível gerar o Word. A seção '{secao}' não foi encontrada."
        
        if "Texto-base" in texto_bruto and "Fonte" not in texto_bruto:
            return False, "Não foi possível gerar o Word. A seção 'Fonte' não foi encontrada (obrigatória para Texto-base)."
            
        return True, ""

    def gerar_docx(self, texto_bruto: str, nome_arquivo: str) -> str:
        valido, erro = self.validar_estrutura(texto_bruto)
        if not valido:
            return erro

        if not self.output_dir.exists():
            self.output_dir.mkdir(parents=True, exist_ok=True)

        caminho_saida = self.output_dir / nome_arquivo
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        linhas = texto_bruto.split('\n')
        for linha in linhas:
            linha = linha.strip()
            if not linha:
                doc.add_paragraph("")
                continue

            p = doc.add_paragraph()
            # Negrito para títulos de seção e alternativas
            is_header = any(linha.startswith(s) for s in self.sections_obrigatorias) or ":" in linha[:20]
            is_alt = any(linha.startswith(f"{l})") for l in "ABCD")
            
            run = p.add_run(linha)
            if is_header or is_alt or "Gabarito" in linha:
                run.bold = True

        doc.save(str(caminho_saida))
        return str(caminho_saida)

    def abrir_documento(self, caminho_arquivo: str) -> str:
        try:
            path = Path(caminho_arquivo)
            if not path.exists():
                return f"Arquivo não encontrado: {caminho_arquivo}"
            
            if sys.platform == "win32":
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.run(["open", str(path)], check=True)
            else:
                subprocess.run(["xdg-open", str(path)], check=True)
            return "Documento aberto com sucesso."
        except Exception:
            return "Documento gerado, mas não foi possível abri-lo automaticamente."
